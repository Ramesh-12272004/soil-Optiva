import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import pandas as pd
import json
import os
import urllib.parse
import base64

# --------------------------
# PAGE CONFIG — must be first
# --------------------------
st.set_page_config(
    page_title="ANITS Soil Lab",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --------------------------
# USER DATABASE
# --------------------------
USER_DB_FILE = "users.json"

def load_users():
    if os.path.exists(USER_DB_FILE):
        with open(USER_DB_FILE, "r") as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(USER_DB_FILE, "w") as f:
        json.dump(users, f)

def register_user(name, email, password):
    users = load_users()
    if email in users:
        return False, "An account with this email already exists."
    users[email] = {"name": name, "password": password}
    save_users(users)
    return True, "Account created successfully!"

def login_user(email, password):
    users = load_users()
    if email not in users:
        return False, None, "No account found with this email."
    if users[email]["password"] != password:
        return False, None, "Incorrect password."
    return True, users[email]["name"], "Login successful!"

# --------------------------
# LOGO HELPER
# --------------------------
def _logo_b64():
    try:
        with open("assets/anits_logo.png", "rb") as f:
            return base64.b64encode(f.read()).decode()
    except Exception:
        return None

_LOGO = _logo_b64()

def logo_img_tag(size=110, radius=20):
    if _LOGO:
        return (
            f'<img src="data:image/png;base64,{_LOGO}" '
            f'style="width:{size}px;height:{size}px;border-radius:{radius}px;'
            f'object-fit:contain;background:white;padding:6px;'
            f'border:2px solid rgba(0,160,255,0.4);'
            f'box-shadow:0 0 40px rgba(0,120,255,0.35);'
            f'display:block;margin:0 auto 18px auto;"/>'
        )
    return f'<div style="font-size:{int(size*0.35)}px;text-align:center;margin-bottom:18px;">🏛️</div>'

# --------------------------
# SESSION STATE
# --------------------------
_defaults = {
    "app_started": False,
    "auth_screen": "login",
    "logged_in": False,
    "user_name": "",
    "user_email": "",
    "completed_tests": {},
    "view_mode": "test",
    "last_result": None,
    "last_test_name": None,
    "page_history": [],
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# --------------------------
# BASE CSS — injected once, every screen
# --------------------------
def inject_base_css():
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700;800;900&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"], .stApp, .stMarkdown,
button, input, textarea, select {
    font-family: 'Sora', sans-serif !important;
}

.stApp, .main, [data-testid="stAppViewContainer"] {
    background-color: #03050f !important;
    background-image:
        radial-gradient(ellipse 60% 40% at 20% 20%, rgba(0,100,255,0.10) 0%, transparent 60%),
        radial-gradient(ellipse 40% 60% at 80% 80%, rgba(0,200,180,0.07) 0%, transparent 60%),
        radial-gradient(ellipse 50% 50% at 50% 50%, rgba(20,60,120,0.12) 0%, transparent 70%) !important;
}

.block-container {
    padding-top: 1rem !important;
    padding-bottom: 4rem !important;
    max-width: 1200px !important;
}

#MainMenu, footer, header,
[data-testid="stToolbar"],
[data-testid="manage-app-button"],
.stDeployButton { display: none !important; }

/* SIDEBAR ALWAYS OPEN */
[data-testid="stSidebar"] {
    background-color: rgba(4,10,28,0.98) !important;
    border-right: 1px solid rgba(0,100,200,0.25) !important;
    min-width: 260px !important;
    max-width: 260px !important;
    transform: translateX(0) !important;
    visibility: visible !important;
    display: flex !important;
    margin-left: 0 !important;
    left: 0 !important;
    position: relative !important;
}
[data-testid="stSidebar"][aria-expanded="false"] {
    min-width: 260px !important;
    max-width: 260px !important;
    transform: translateX(0) !important;
    visibility: visible !important;
    display: flex !important;
    margin-left: 0 !important;
}
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"],
button[data-testid="baseButton-header"] {
    display: none !important;
}

[data-testid="stSidebar"] * { color: rgba(210,232,255,0.92) !important; }

.stTextInput input, .stNumberInput input, .stTextArea textarea {
    background: rgba(0,20,60,0.7) !important;
    border: 1px solid rgba(0,100,200,0.45) !important;
    border-radius: 10px !important;
    color: #e8f4ff !important;
    padding: 0.6rem 1rem !important;
}
.stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus {
    border-color: rgba(0,190,255,0.7) !important;
    box-shadow: 0 0 0 3px rgba(0,150,255,0.15) !important;
    outline: none !important;
}
.stTextInput label, .stNumberInput label,
.stSelectbox label, .stTextArea label,
.stRadio label {
    color: rgba(200,228,255,0.95) !important;
    font-weight: 600 !important;
    font-size: 0.87rem !important;
}

.stButton > button {
    background: linear-gradient(135deg, #0064ff, #003cc8) !important;
    color: #ffffff !important;
    border: 1px solid rgba(0,160,255,0.5) !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    padding: 0.6rem 1.4rem !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 14px rgba(0,80,200,0.3) !important;
    width: 100% !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #0088ff, #0050e0) !important;
    box-shadow: 0 8px 24px rgba(0,130,255,0.45) !important;
    transform: translateY(-2px) !important;
}

.stDownloadButton > button {
    background: linear-gradient(135deg, #00a878, #006650) !important;
    border-color: rgba(0,200,140,0.5) !important;
    color: #fff !important;
}

[data-testid="stLinkButton"] a {
    background: linear-gradient(135deg, rgba(0,100,255,0.8), rgba(0,60,185,0.9)) !important;
    color: #fff !important;
    border: 1px solid rgba(0,140,255,0.4) !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    text-decoration: none !important;
    display: block !important;
    transition: all 0.2s ease !important;
}
[data-testid="stLinkButton"] a:hover {
    transform: translateY(-3px) !important;
    box-shadow: 0 8px 22px rgba(0,140,255,0.4) !important;
}

.stSelectbox > div > div {
    background: rgba(0,20,60,0.7) !important;
    border: 1px solid rgba(0,100,200,0.4) !important;
    border-radius: 10px !important;
    color: #e8f4ff !important;
}

[data-testid="stExpander"] {
    background: rgba(0,20,60,0.4) !important;
    border: 1px solid rgba(0,100,200,0.3) !important;
    border-radius: 12px !important;
}
[data-testid="stExpander"] summary {
    color: rgba(205,228,255,0.95) !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    padding: 10px 14px !important;
    background: rgba(0,40,100,0.4) !important;
    border-radius: 12px !important;
    list-style: none !important;
}
[data-testid="stExpander"] summary::-webkit-details-marker { display: none !important; }
[data-testid="stExpander"] summary p {
    display: inline !important;
    color: rgba(205,228,255,0.95) !important;
    font-weight: 600 !important;
}

.stTabs [data-baseweb="tab"] { color: rgba(180,215,255,0.75) !important; font-weight: 600 !important; }
.stTabs [aria-selected="true"] { color: #fff !important; border-bottom: 2px solid #0088ff !important; }

[data-testid="stDataFrame"] {
    background: rgba(0,20,60,0.4) !important;
    border: 1px solid rgba(0,100,200,0.25) !important;
    border-radius: 10px !important;
}

[data-testid="stChatMessage"] {
    background: rgba(0,20,60,0.45) !important;
    border: 1px solid rgba(0,100,200,0.22) !important;
    border-radius: 12px !important;
}

.stMarkdown p, .stMarkdown li { color: rgba(205,228,255,0.9) !important; }
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 { color: #ffffff !important; }
p { color: rgba(205,228,255,0.88) !important; }
strong, b { color: #ffffff !important; }

::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: rgba(0,20,60,0.3); }
::-webkit-scrollbar-thumb { background: rgba(0,100,200,0.5); border-radius: 3px; }

.stNumberInput button {
    background: rgba(0,40,120,0.5) !important;
    border: 1px solid rgba(0,100,200,0.3) !important;
    color: rgba(180,220,255,0.9) !important;
}

[data-testid="stAlert"] p { color: #fff !important; }

.greet-bar {
    background: linear-gradient(135deg, rgba(0,100,255,0.4), rgba(0,180,200,0.3));
    border: 1px solid rgba(0,140,255,0.45);
    color: #fff !important;
    padding: 11px 14px;
    border-radius: 11px;
    font-weight: 700;
    font-size: 0.9rem;
    text-align: center;
    margin-bottom: 8px;
}

.header-banner {
    background: linear-gradient(135deg, rgba(0,55,175,0.7), rgba(0,115,195,0.5), rgba(0,75,160,0.7));
    border: 1px solid rgba(0,150,255,0.3);
    padding: 18px 28px;
    border-radius: 16px;
    text-align: center;
    margin-bottom: 24px;
    box-shadow: 0 8px 32px rgba(0,0,0,0.4);
}
.header-banner h2 { font-size: clamp(1.1rem,2.5vw,1.5rem) !important; font-weight: 800 !important; margin: 0 0 4px 0 !important; color: #fff !important; }
.header-banner p  { margin: 0 !important; color: rgba(195,228,255,0.92) !important; font-size: 0.83rem !important; text-transform: uppercase !important; letter-spacing: 0.06em !important; }

.metric-row { display: flex; gap: 12px; margin-bottom: 20px; flex-wrap: wrap; }
.metric-card { background: rgba(0,40,120,0.35); border: 1px solid rgba(0,120,255,0.25); border-radius: 13px; padding: 14px 20px; flex: 1; min-width: 120px; }
.metric-card:hover { background: rgba(0,60,180,0.45); transform: translateY(-4px); box-shadow: 0 8px 28px rgba(0,100,255,0.2); transition: all 0.25s ease; }
.metric-label { font-size: 0.7rem; color: rgba(160,205,255,0.85); font-weight: 700; text-transform: uppercase; letter-spacing: 0.1em; margin-bottom: 3px; }
.metric-value { font-size: 1.3rem; font-weight: 800; color: #fff; font-family: 'JetBrains Mono', monospace !important; }

.rec-card { border-radius: 11px; padding: 13px 17px; margin-bottom: 8px; }
.rec-card.danger  { background: rgba(200,40,40,0.12);  border: 1px solid rgba(200,40,40,0.28);  border-left: 4px solid #ff4444; }
.rec-card.warning { background: rgba(200,150,0,0.12);  border: 1px solid rgba(200,150,0,0.28);  border-left: 4px solid #ffaa00; }
.rec-card.success { background: rgba(0,180,100,0.12);  border: 1px solid rgba(0,180,100,0.28);  border-left: 4px solid #00cc66; }
.rec-card.info    { background: rgba(0,120,220,0.12);  border: 1px solid rgba(0,120,220,0.28);  border-left: 4px solid #4488ff; }
.rec-title { font-size: 0.9rem; font-weight: 800; color: #fff; margin-bottom: 4px; }
.rec-body  { font-size: 0.81rem; color: rgba(210,232,255,0.85); line-height: 1.55; }
.rec-section-title { font-size: 0.95rem; font-weight: 800; color: rgba(100,210,255,0.95); letter-spacing: 0.05em; text-transform: uppercase; margin-bottom: 10px; padding-bottom: 6px; border-bottom: 1px solid rgba(0,120,200,0.3); }

.soil-badge { display: inline-flex; align-items: center; gap: 12px; background: linear-gradient(135deg, rgba(0,60,180,0.5), rgba(0,120,200,0.35)); border: 1px solid rgba(0,160,255,0.45); border-radius: 100px; padding: 11px 24px; margin: 10px 0; }
.soil-badge-symbol { font-size: 1.45rem; font-weight: 900; color: #00d4ff; font-family: 'JetBrains Mono', monospace !important; }
.soil-badge-name   { font-size: 0.9rem; font-weight: 600; color: rgba(200,235,255,0.95); }

.hist-card { background: rgba(0,20,60,0.55); border: 1px solid rgba(0,100,200,0.22); border-radius: 13px; padding: 16px 20px 12px; margin-bottom: 12px; border-left: 4px solid rgba(0,140,255,0.65); }
.hist-test-name { font-size: 0.96rem; font-weight: 800; color: #fff; margin-bottom: 2px; }
.hist-time { font-size: 0.74rem; color: rgba(160,200,245,0.75); margin-bottom: 8px; }
.hist-prop { display: inline-block; background: rgba(0,80,200,0.3); border: 1px solid rgba(0,120,255,0.3); color: rgba(200,230,255,0.95); border-radius: 5px; padding: 2px 9px; font-size: 0.74rem; font-weight: 700; margin: 2px 3px 2px 0; font-family: 'JetBrains Mono', monospace !important; }

.class-section { background: rgba(0,40,100,0.38); border: 1px solid rgba(0,140,255,0.28); border-radius: 13px; padding: 18px 22px; margin-bottom: 14px; }
.share-section { background: rgba(0,30,80,0.55); border: 1px solid rgba(0,120,200,0.28); border-radius: 14px; padding: 22px 26px; margin-top: 24px; border-top: 3px solid rgba(0,185,255,0.45); }

.feat-grid { display: grid; grid-template-columns: repeat(3,1fr); gap: 8px; margin: 0 auto 40px auto; width: 100%; max-width: 580px; }
.feat-item { background: rgba(0,60,160,0.28); border: 1px solid rgba(0,120,255,0.28); border-radius: 10px; padding: 10px 14px; font-size: 0.78rem; color: rgba(200,230,255,0.92); font-weight: 600; display: flex; align-items: center; gap: 6px; }

.launch-wrap .stButton > button {
    font-size: 1.12rem !important;
    padding: 0.82rem 0 !important;
    border-radius: 50px !important;
    border: 2px solid rgba(0,210,255,0.55) !important;
    box-shadow: 0 6px 30px rgba(0,120,255,0.5) !important;
    min-height: 54px !important;
}

@media (max-width: 768px) {
    .feat-grid { grid-template-columns: repeat(2,1fr) !important; }
    .metric-row { flex-direction: column !important; }
    .block-container { padding-left: 0.8rem !important; padding-right: 0.8rem !important; }
}
</style>
""", unsafe_allow_html=True)

inject_base_css()

# --------------------------
# AI CHATBOT
# --------------------------
def get_ai_response(query):
    q = query.lower().strip()
    if "cbr" in q and "low" in q:
        return "🔴 Low CBR means weak subgrade soil. Soil stabilisation or thicker pavement required (IRC:37). CBR < 3% not suitable for direct use."
    if "cbr" in q:
        return "📏 CBR measures soil strength for road subgrade design. CBR > 15% is ideal; < 3% needs stabilisation (IS 2720 Part 16)."
    if "liquid limit" in q or "ll" in q:
        return "💧 Liquid Limit: water content where soil transitions from plastic to liquid. LL > 50% = high compressibility (IS 2720 Part 5)."
    if "plastic limit" in q or "pl" in q:
        return "🧱 Plastic Limit: minimum water content at which soil rolls into 3mm thread without crumbling."
    if "plasticity index" in q or "pi" in q:
        return "📊 PI = LL − PL. PI < 7 = low plasticity. PI > 17 = high plasticity clay — shrink-swell risk (IS 1904)."
    if "atterberg" in q:
        return "📋 Atterberg Limits: LL, PL, SL — define soil consistency states. Key for classification (IS 1498)."
    if "shrinkage limit" in q:
        return "🔵 Shrinkage Limit: water content below which soil volume stops reducing on drying."
    if "proctor" in q or "compaction" in q or "omc" in q or "mdd" in q:
        return "⚙️ Proctor test finds OMC and MDD. MDD > 1.9 g/cc = suitable for embankments (IS 2720 Part 7)."
    if "shear strength" in q or "direct shear" in q:
        return "💪 Shear strength = c + σ·tan(φ). Via Direct Shear, Triaxial, or Vane Shear (IS 2720 Parts 12–13)."
    if "triaxial" in q:
        return "🔬 Triaxial test: shear strength under controlled drainage. Types: UU, CU, CD."
    if "ucs" in q or "unconfined" in q:
        return "📐 UCS = 2c for saturated clays. qu < 25 kPa = very soft; > 100 kPa = stiff (IS 6403)."
    if "vane shear" in q:
        return "🌀 Vane Shear: undrained shear strength in soft cohesive soils, often in-situ."
    if "permeability" in q or "constant head" in q or "variable head" in q:
        return "💧 Constant head for sandy soils; Falling head for fine-grained soils (IS 2720 Part 17)."
    if "consolidation" in q or "settlement" in q:
        return "📉 Consolidation test gives Cv and Cc. Low Cv = slow, large settlement (IS 2720 Part 15)."
    if "sieve" in q or "grain size" in q or "gradation" in q:
        return "🔎 Sieve Analysis classifies soil by particle size. Cu and Cc determine gradation quality (IS 2720 Part 4)."
    if "specific gravity" in q:
        return "⚖️ Gs typically 2.65–2.80. Used to compute void ratio and degree of saturation (IS 2720 Part 3)."
    if "core cutter" in q or "bulk density" in q or "field density" in q:
        return "🔩 Core Cutter: in-situ bulk density. Quick method for cohesive soils (IS 2720 Part 29)."
    if "uscs" in q or "classification" in q or "is 1498" in q:
        return "📚 IS 1498/USCS: GW, GP, GM, GC, SW, SP, SM, SC, ML, CL, MI, CI, MH, CH etc."
    if "expansive" in q or "black cotton" in q:
        return "⚠️ FSI > 50% = expansive soil. Use lime stabilisation or under-reamed piles."
    if "foundation" in q:
        return "🏗️ Soft clay → Raft/Pile. Medium → Isolated footings. Always check IS 1904 & IS 6403."
    if "bearing capacity" in q:
        return "📐 Use Terzaghi/Meyerhof equations. FOS ≥ 3 always (IS 6403)."
    if "pile" in q:
        return "🔩 Friction piles: load via skin friction. End-bearing piles: rest on hard stratum."
    if "stabiliz" in q or "lime" in q or "cement" in q:
        return "🛠️ Lime reduces plasticity; cement increases strength. Also fly ash/geotextiles (IRC:SP:20)."
    if "help" in q or "how to" in q:
        return "📱 Select test from sidebar → enter observations → click Calculate. DOCX auto-generated."
    if "report" in q or "download" in q or "docx" in q:
        return "📥 Every test generates a downloadable DOCX with procedure, formulas, results, IS Code recs."
    if "is 2720" in q:
        return "📖 IS 2720: Methods of Test for Soils — 40+ parts (classification, compaction, shear, consolidation)."
    if "is code" in q or "irc" in q:
        return "📖 IS 1498 (Classification), IS 2720 (Tests), IS 1904 (Foundation), IS 6403 (Bearing Capacity), IRC:37 (Pavement)."
    return "🤖 Ask me about soil tests, IS codes, or foundation design. Try: 'What is CBR?', 'Explain liquid limit', 'Foundation for soft clay?'"

def ai_chatbot(key_prefix="main"):
    chat_key  = f"chat_history_{key_prefix}"
    input_key = f"ai_input_{key_prefix}"
    ask_key   = f"ask_btn_{key_prefix}"
    if chat_key not in st.session_state:
        st.session_state[chat_key] = [("Bot", "👋 Hi! I'm your Soil Testing AI Assistant. Ask me anything about soil tests, IS codes, or foundation design!")]
    st.markdown("### 🤖 AI Assistant")
    st.markdown("<p style='color:rgba(190,220,255,0.9);font-weight:600;font-size:0.85rem;'>💡 Quick Questions</p>", unsafe_allow_html=True)
    qcols = st.columns(2)
    quick_prompts = [
        ("What is CBR?", "What is CBR?"),
        ("Explain Liquid Limit", "Explain liquid limit"),
        ("Best foundation for soft clay?", "Best foundation for soft clay?"),
        ("What is shear strength?", "What is shear strength?"),
    ]
    for i, (label, prompt) in enumerate(quick_prompts):
        with qcols[i % 2]:
            if st.button(label, key=f"qp_{key_prefix}_{i}", width='stretch'):
                st.session_state[chat_key].append(("You", label))
                st.session_state[chat_key].append(("Bot", get_ai_response(prompt)))
                st.rerun()
    for role, msg in st.session_state[chat_key][-8:]:
        with st.chat_message("assistant" if role == "Bot" else "user"):
            st.markdown(msg)
    col_in, col_btn = st.columns([4, 1])
    with col_in:
        user_input = st.text_input("Ask a question", placeholder="Ask about soil tests, IS codes…", key=input_key, label_visibility="collapsed")
    with col_btn:
        if st.button("Ask →", key=ask_key, width='stretch'):
            if user_input and user_input.strip():
                st.session_state[chat_key].append(("You", user_input.strip()))
                st.session_state[chat_key].append(("Bot", get_ai_response(user_input)))
                st.rerun()

# --------------------------
# IS CODE RECOMMENDATIONS
# --------------------------
def get_is_recommendations(test_name, result_dict):
    recs = []
    vals = {k.lower().replace(" ", "_"): v for k, v in result_dict.items() if isinstance(v, (int, float))}
    if "cbr" in test_name.lower():
        cbr = vals.get("cbr_%", vals.get("cbr_value", None))
        if cbr is not None:
            if cbr < 3:    recs.append(("🔴 CBR < 3%",  "Not suitable for subgrade. Stabilisation required (IS 2720 Part 16).", "danger"))
            elif cbr < 7:  recs.append(("🟡 CBR 3–7%",  "Weak subgrade. Thick pavement needed (IRC:37).", "warning"))
            elif cbr < 15: recs.append(("🟢 CBR 7–15%", "Moderate subgrade. Suitable with appropriate pavement.", "success"))
            else:           recs.append(("✅ CBR > 15%", "Good subgrade. Economical pavement possible (IRC:37).", "success"))
    if "liquid limit" in test_name.lower() or "ll" in vals:
        ll = vals.get("liquid_limit_%", vals.get("ll", None))
        if ll is not None:
            if ll < 35:    recs.append(("🟢 LL < 35%",  "Low plasticity (ML/CL). Good for earthworks (IS 1498).", "success"))
            elif ll < 50:  recs.append(("🟡 LL 35–50%", "Medium plasticity. Use caution in foundations.", "warning"))
            else:           recs.append(("🔴 LL > 50%",  "High compressibility. Avoid direct foundation (IS 1904).", "danger"))
    if "plastic" in test_name.lower():
        pi = vals.get("plasticity_index", vals.get("pi", None))
        if pi is not None:
            if pi < 7:    recs.append(("✅ PI < 7",   "Low plasticity. Suitable for pavement subgrade.", "success"))
            elif pi < 17: recs.append(("🟡 PI 7–17",  "Medium plasticity. Monitor swelling.", "warning"))
            else:          recs.append(("🔴 PI > 17",  "High plasticity — shrink-swell risk (IS 1904).", "danger"))
    if "compaction" in test_name.lower():
        mdd = vals.get("mdd_g/cc", vals.get("maximum_dry_density", None))
        if mdd is not None:
            if mdd > 1.9:   recs.append(("✅ MDD > 1.9 g/cc",    "Dense soil. Good for embankment (IS 2720 Part 7).", "success"))
            elif mdd > 1.6: recs.append(("🟡 MDD 1.6–1.9 g/cc", "Moderate density. Suitable for earthwork.", "warning"))
            else:            recs.append(("🔴 MDD < 1.6 g/cc",   "Low density. Compaction improvement needed.", "danger"))
    if any(x in test_name.lower() for x in ["shear", "ucs", "triaxial"]):
        qu = vals.get("qu_kn/m²", vals.get("unconfined_compressive_strength", None))
        if qu is not None:
            if qu < 25:    recs.append(("🔴 qu < 25 kPa",   "Very soft clay. Not for direct loading (IS 6403).", "danger"))
            elif qu < 100: recs.append(("🟡 qu 25–100 kPa", "Soft–medium clay. Bearing capacity check needed.", "warning"))
            else:           recs.append(("✅ qu > 100 kPa",  "Stiff clay. Good for light structures.", "success"))
    if "consolidation" in test_name.lower():
        cv = vals.get("cv_cm²/s", vals.get("coefficient_of_consolidation", None))
        if cv is not None:
            if cv < 0.001: recs.append(("🔴 Low Cv",      "Very slow consolidation. Large settlements expected (IS 2720 Part 15).", "danger"))
            else:           recs.append(("🟢 Adequate Cv", "Consolidation rate acceptable.", "success"))
    if not recs:
        recs.append(("ℹ️ No specific IS recommendation", "Manual interpretation required. Refer IS 2720 series.", "info"))
    return recs

def get_soil_classification(result_dict):
    vals = {k.lower().replace(" ", "_"): v for k, v in result_dict.items() if isinstance(v, (int, float))}
    ll = vals.get("liquid_limit_%", vals.get("ll", None))
    pi = vals.get("plasticity_index", vals.get("pi", None))
    if ll is not None and pi is not None:
        if ll < 35:   return ("ML", "Silt of low plasticity", "🟡") if pi < 7 else ("CL", "Clay of low plasticity", "🟢")
        elif ll < 50: return ("MI", "Silt of intermediate plasticity", "🟡") if pi < 7 else ("CI", "Clay of intermediate plasticity", "🟠")
        else:          return ("MH", "Silt of high plasticity", "🔴") if pi < 7 else ("CH", "Clay of high plasticity", "🔴")
    return None, None, None

# --------------------------
# SHARING
# --------------------------
def build_share_text(test_name, result_dict):
    lines = [f"🧪 *ANITS Soil Test Report*", f"Test: *{test_name}*", ""]
    for k, v in result_dict.items():
        if isinstance(v, (int, float)): lines.append(f"• {k}: {round(v, 3)}")
        elif isinstance(v, str) and k not in ("procedure", "formulas") and len(v) < 120: lines.append(f"• {k}: {v}")
    lines.append("\n_Generated by ANITS Civil Dept – Soil Testing System_")
    return "\n".join(lines)

def build_ai_prompt(test_name, result_dict):
    lines = [f"Soil test results from '{test_name}' at ANITS Civil Engineering Laboratory.\n\nTest Results:"]
    for k, v in result_dict.items():
        if isinstance(v, (int, float)): lines.append(f"  - {k}: {round(v, 3)}")
        elif isinstance(v, str) and k not in ("procedure", "formulas", "data", "graph", "diagram") and len(v) < 120: lines.append(f"  - {k}: {v}")
    lines += ["\nPlease provide: 1. Result interpretation  2. IS Code references  3. Foundation/pavement recommendations  4. Stabilisation suggestions if needed"]
    return "\n".join(lines)

def share_buttons(test_name, result_dict, doc_bytes=None, inside_expander=False):
    text     = build_share_text(test_name, result_dict)
    ai_p     = build_ai_prompt(test_name, result_dict)
    enc_t    = urllib.parse.quote(text)
    enc_ai   = urllib.parse.quote(ai_p)
    search_q = urllib.parse.quote(f"{test_name} IS code soil test India")

    st.markdown("#### 📤 Share Results")
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.link_button("💬 WhatsApp",  f"https://api.whatsapp.com/send?text={enc_t}", width='stretch')
    with c2: st.link_button("✈️ Telegram",  f"https://t.me/share/url?url=&text={enc_t}", width='stretch')
    with c3: st.link_button("📧 Email",     f"mailto:?subject={urllib.parse.quote('ANITS Report – ' + test_name)}&body={urllib.parse.quote(text.replace('*', ''))}", width='stretch')
    with c4: st.link_button("🐦 Twitter/X", f"https://twitter.com/intent/tweet?text={enc_t}", width='stretch')

    st.markdown("#### 🔍 Search & AI Assistants")
    cg, cs, cgpt, ccop = st.columns(4)
    with cg:   st.link_button("🔍 Google",  f"https://www.google.com/search?q={search_q}", width='stretch')
    with cs:   st.link_button("📚 Scholar", f"https://scholar.google.com/scholar?q={search_q}", width='stretch')
    with cgpt: st.link_button("🟢 ChatGPT", f"https://chat.openai.com/?q={enc_ai}", width='stretch')
    with ccop: st.link_button("🔵 Copilot", "https://copilot.microsoft.com/", width='stretch')

    st.text_area("📋 Copilot Prompt — copy & paste:", value=ai_p, height=120,
                 key=f"cpta_{abs(hash(test_name + str(id(result_dict))))}", label_visibility="visible")
    if doc_bytes:
        st.download_button(
            "📥 Download Report (.docx)",
            data=doc_bytes,
            file_name=f"ANITS_{test_name.replace(' ', '_')}_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width='stretch',
        )

# --------------------------
# DOCX BUILDERS
# --------------------------
def _fill_doc_for_test(doc, test_name, res):
    doc.add_heading(test_name, 1)
    for section, heading in [("procedure", "Procedure"), ("formulas", "Formulas")]:
        if section in res:
            doc.add_heading(heading, 2)
            for line in str(res[section]).split("\n"):
                if line.strip(): doc.add_paragraph(line.strip())
    if "data" in res and isinstance(res["data"], pd.DataFrame):
        doc.add_heading("Results Data", 2)
        df  = res["data"]
        tbl = doc.add_table(rows=1, cols=len(df.columns))
        tbl.style = "Table Grid"
        for i, col in enumerate(df.columns): tbl.rows[0].cells[i].text = str(col)
        for _, row in df.iterrows():
            cells = tbl.add_row().cells
            for i, val in enumerate(row): cells[i].text = str(val)
    for img_key, heading in [("graph", "Graph"), ("diagram", "Diagram")]:
        if img_key in res and res[img_key] is not None:
            doc.add_heading(heading, 2)
            try:
                res[img_key].seek(0)
                doc.add_picture(res[img_key], width=Inches(5 if img_key == "graph" else 4))
            except Exception:
                pass
    doc.add_heading("IS Code Recommendations", 2)
    for title, body, _ in get_is_recommendations(test_name, res):
        doc.add_paragraph(f"{title}: {body}")
    for key, value in res.items():
        if isinstance(value, str) and key not in ["procedure", "formulas", "data", "graph", "diagram"]:
            doc.add_paragraph(f"{key}: {value}")

def build_single_test_docx(test_name, res):
    doc = Document()
    doc.add_heading(f"ANITS Soil Test Report – {test_name}", 0)
    _fill_doc_for_test(doc, test_name, res)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_all_tests_docx(completed_tests):
    doc = Document()
    doc.add_heading("ANITS Soil Test Report – All Tests", 0)
    first = True
    for name, res in completed_tests.items():
        if not first: doc.add_page_break()
        first = False
        _fill_doc_for_test(doc, name, res)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# --------------------------
# BACK BUTTON
# --------------------------
def show_back_button():
    if st.session_state.page_history:
        _, col, _ = st.columns([1, 2, 1])
        with col:
            if st.button("⬅️  Back to Previous Page", key="back_btn_main", width='stretch'):
                st.session_state.view_mode = st.session_state.page_history.pop()
                st.rerun()

# ==========================
# SCREEN 1 — WELCOME
# ==========================
if not st.session_state.app_started:
    st.markdown(f"""
    <div style="display:flex;flex-direction:column;align-items:center;
        justify-content:center;text-align:center;padding:48px 20px 20px;">
        {logo_img_tag(size=118, radius=22)}
        <div style="display:inline-flex;align-items:center;gap:8px;
            background:rgba(0,100,255,0.15);border:1px solid rgba(0,160,255,0.38);
            border-radius:100px;padding:6px 20px;font-size:0.76rem;font-weight:700;
            color:rgba(200,230,255,0.95);letter-spacing:0.1em;text-transform:uppercase;margin-bottom:18px;">
            🏗️ ANITS · Civil Engineering
        </div>
        <div style="font-size:clamp(2rem,5vw,3.2rem);font-weight:900;color:#fff;
            line-height:1.15;margin-bottom:14px;letter-spacing:-0.02em;">
            Soil Testing<br>
            <span style="background:linear-gradient(135deg,#0099ff,#00ddbb);
                -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">
                Analysis System</span>
        </div>
        <div style="font-size:0.95rem;color:rgba(185,218,255,0.85);line-height:1.65;
            margin-bottom:28px;max-width:500px;">
            Professional soil lab platform with IS Code recommendations,
            AI classification, auto-generated reports, and multi-platform sharing.
        </div>
        <div class="feat-grid">
            <div class="feat-item">🧪 15 IS Standard Tests</div>
            <div class="feat-item">📋 IS Code Recommendations</div>
            <div class="feat-item">🤖 AI Soil Classification</div>
            <div class="feat-item">📥 Auto DOCX Reports</div>
            <div class="feat-item">📊 Charts &amp; Graphs</div>
            <div class="feat-item">📤 Share via WhatsApp</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    _, mid, _ = st.columns([1.8, 1, 1.8])
    with mid:
        st.markdown('<div class="launch-wrap">', unsafe_allow_html=True)
        if st.button("🚀  Launch App →", key="launch_btn", width='stretch'):
            st.session_state.app_started = True
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    st.stop()

# ==========================
# SCREEN 2 — AUTH
# ==========================
elif not st.session_state.logged_in:
    st.markdown("<div style='height:24px'></div>", unsafe_allow_html=True)
    _, mid, _ = st.columns([1, 2, 1])
    with mid:
        if _LOGO:
            st.markdown(
                f'<div style="text-align:center;margin-bottom:12px;">'
                f'<img src="data:image/png;base64,{_LOGO}" style="width:88px;height:88px;'
                f'border-radius:16px;object-fit:contain;background:white;padding:6px;'
                f'border:2px solid rgba(0,160,255,0.38);box-shadow:0 0 28px rgba(0,120,255,0.32);"/>'
                f'</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div style="text-align:center;font-size:3.5rem;margin-bottom:12px;">🏛️</div>', unsafe_allow_html=True)

        if st.session_state.auth_screen == "login":
            st.markdown('<h2 style="text-align:center;color:#fff;margin-bottom:4px;">Welcome Back 👋</h2>', unsafe_allow_html=True)
            st.markdown('<p style="text-align:center;color:rgba(175,215,255,0.8);margin-bottom:20px;font-size:0.84rem;">ANITS · Soil Testing System</p>', unsafe_allow_html=True)
            email    = st.text_input("📧  Email address", key="login_email", placeholder="you@example.com")
            password = st.text_input("🔒  Password", type="password", key="login_pass", placeholder="Enter your password")
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            if st.button("🔐  Sign In", width='stretch', key="signin_btn"):
                if not email or not password:
                    st.error("Please fill in all fields.")
                else:
                    ok, name, msg = login_user(email.strip().lower(), password)
                    if ok:
                        st.session_state.logged_in  = True
                        st.session_state.user_name  = name
                        st.session_state.user_email = email.strip().lower()
                        st.session_state.view_mode  = "test"
                        st.rerun()
                    else:
                        st.error(msg)
            st.markdown("<div style='text-align:center;color:rgba(150,190,230,0.55);font-size:0.8rem;margin:12px 0;'>── or ──</div>", unsafe_allow_html=True)
            if st.button("✏️  Create a New Account", width='stretch', key="goto_signup"):
                st.session_state.auth_screen = "signup"
                st.rerun()
        else:
            st.markdown('<h2 style="text-align:center;color:#fff;margin-bottom:4px;">Create Account 🎓</h2>', unsafe_allow_html=True)
            st.markdown('<p style="text-align:center;color:rgba(175,215,255,0.8);margin-bottom:20px;font-size:0.84rem;">Join ANITS Soil Testing System</p>', unsafe_allow_html=True)
            full_name = st.text_input("👤  Full Name",        key="reg_name",  placeholder="e.g. Ravi Kumar")
            email     = st.text_input("📧  Email Address",    key="reg_email", placeholder="you@example.com")
            password  = st.text_input("🔒  Password",         type="password", key="reg_pass",  placeholder="Min. 6 characters")
            confirm   = st.text_input("🔒  Confirm Password", type="password", key="reg_pass2", placeholder="Re-enter password")
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            if st.button("🎓  Create Account", width='stretch', key="create_btn"):
                if not all([full_name, email, password, confirm]):
                    st.error("Please fill in all fields.")
                elif len(password) < 6:
                    st.error("Password must be at least 6 characters.")
                elif password != confirm:
                    st.error("Passwords do not match.")
                else:
                    ok, msg = register_user(full_name.strip(), email.strip().lower(), password)
                    if ok:
                        st.success(msg + " Please sign in.")
                        st.session_state.auth_screen = "login"
                        st.rerun()
                    else:
                        st.error(msg)
            st.markdown("<div style='text-align:center;color:rgba(150,190,230,0.55);font-size:0.8rem;margin:12px 0;'>── or ──</div>", unsafe_allow_html=True)
            if st.button("🔐  Sign In Instead", width='stretch', key="goto_login"):
                st.session_state.auth_screen = "login"
                st.rerun()

    st.stop()

# ==========================
# SCREEN 3 — MAIN APP
# ==========================
else:
    try:
        from history_manager import load_history, save_history, clear_history
    except ImportError:
        def load_history(email): return []
        def save_history(email, name, result): pass
        def clear_history(email): pass

    try:
        from tabs import (
            sieve_analysis, liquid_limit_casagrande, liquid_limit_cone,
            plastic_limit, core_cutter, specific_gravity, constant_head,
            variable_head, compaction_test, direct_shear, ucs_test,
            consolidation, cbr_test, vane_shear, triaxial_test,
        )
        tests = {
            "Sieve Analysis":            sieve_analysis,
            "Liquid Limit (Casagrande)": liquid_limit_casagrande,
            "Liquid Limit (Cone)":       liquid_limit_cone,
            "Plastic Limit":             plastic_limit,
            "Core Cutter":               core_cutter,
            "Specific Gravity":          specific_gravity,
            "Constant Head":             constant_head,
            "Variable Head":             variable_head,
            "Light Compaction":          compaction_test,
            "Direct Shear":              direct_shear,
            "UCS Test":                  ucs_test,
            "Triaxial Test":             triaxial_test,
            "Vane Shear":                vane_shear,
            "CBR Test":                  cbr_test,
            "Consolidation Test":        consolidation,
        }
    except ImportError:
        tests = {"Demo Test": None}

    # ── SIDEBAR ──
    try:
        st.sidebar.image("assets/anits_logo.png", width=90)
    except Exception:
        st.sidebar.markdown('<div style="text-align:center;font-size:2.5rem;">🏛️</div>', unsafe_allow_html=True)

    history_all = load_history(st.session_state.user_email)
    st.sidebar.markdown(f'<div class="greet-bar">👋 Hello, {st.session_state.user_name}!</div>', unsafe_allow_html=True)
    st.sidebar.markdown(
        f"<p style='text-align:center;color:rgba(185,222,255,0.88);font-size:0.82rem;margin-bottom:8px;'>"
        f"📊 <b style='color:#fff;'>{len(history_all)}</b> tests &nbsp;|&nbsp; "
        f"<b style='color:#fff;'>{len(set(e['test_name'] for e in history_all))}</b> types</p>",
        unsafe_allow_html=True)
    st.sidebar.markdown("---")

    _nav_idx = {"test": 0, "history": 1, "ai": 2}.get(st.session_state.view_mode, 0)
    nav = st.sidebar.radio("Navigation", ["🧪  Run Tests", "🕒  Test History", "🤖  AI Assistant"], index=_nav_idx)
    new_mode = "test" if "Run" in nav else ("history" if "History" in nav else "ai")
    if new_mode != st.session_state.view_mode:
        st.session_state.page_history.append(st.session_state.view_mode)
        st.session_state.view_mode = new_mode

    selected_test = selected_module = None
    if st.session_state.view_mode == "test":
        st.sidebar.markdown("---")
        st.sidebar.subheader("Select Test")
        selected_test   = st.sidebar.radio("Select a Test", list(tests.keys()), label_visibility="collapsed")
        selected_module = tests[selected_test]

    st.sidebar.markdown("---")
    if st.session_state.completed_tests:
        st.sidebar.markdown(
            f"<p style='color:rgba(185,222,255,0.9);font-size:0.82rem;margin-bottom:6px;'>"
            f"✅ <b style='color:white;'>{len(st.session_state.completed_tests)}</b> test(s) in session</p>",
            unsafe_allow_html=True)
        st.sidebar.download_button(
            "📥 Download Full Report",
            data=build_all_tests_docx(st.session_state.completed_tests),
            file_name="ANITS_All_Tests_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if st.sidebar.button("🚪  Logout", width='stretch'):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

    # ── HEADER ──
    st.markdown("""
    <div class="header-banner">
        <h2>🏗️ Soil Tests Analysis Dashboard</h2>
        <p>ANITS · Department of Civil Engineering · Professional Lab System</p>
    </div>
    """, unsafe_allow_html=True)

    # ── AI VIEW ──
    if st.session_state.view_mode == "ai":
        show_back_button()
        st.markdown('<h3 style="color:#fff;margin-bottom:4px;">🤖 AI Soil Assistant</h3>', unsafe_allow_html=True)
        st.markdown('<p style="color:rgba(180,215,255,0.82);margin-bottom:16px;font-size:0.85rem;">Ask anything about soil tests, IS codes, or engineering design.</p>', unsafe_allow_html=True)
        ai_chatbot(key_prefix="inapp")

    # ── HISTORY VIEW ──
    elif st.session_state.view_mode == "history":
        show_back_button()
        history = load_history(st.session_state.user_email)
        col_title, col_clear = st.columns([5, 1])
        with col_title:
            st.markdown('<h3 style="color:#fff;margin-bottom:2px;">🕒 Test History</h3>', unsafe_allow_html=True)
            st.markdown('<p style="color:rgba(180,215,255,0.8);font-size:0.85rem;">All previously performed soil tests — newest first.</p>', unsafe_allow_html=True)
        with col_clear:
            if history and st.button("🗑️ Clear All"):
                clear_history(st.session_state.user_email)
                st.success("History cleared.")
                st.rerun()

        if not history:
            st.markdown('<div style="text-align:center;padding:60px 20px;"><div style="font-size:3rem;margin-bottom:12px;">🧪</div><p style="color:rgba(160,200,245,0.7);">No test history yet.</p></div>', unsafe_allow_html=True)
        else:
            unique_types = set(e["test_name"] for e in history)
            st.markdown(f"""
            <div class="metric-row">
                <div class="metric-card"><div class="metric-label">Total Tests</div><div class="metric-value">{len(history)}</div></div>
                <div class="metric-card"><div class="metric-label">Test Types</div><div class="metric-value">{len(unique_types)}</div></div>
                <div class="metric-card"><div class="metric-label">Latest Test</div><div class="metric-value" style="font-size:0.82rem;">{history[0]['test_name']}</div></div>
            </div>""", unsafe_allow_html=True)

            fc = st.selectbox("🔍 Filter by test type", ["All Tests"] + sorted(unique_types))
            filtered = history if fc == "All Tests" else [e for e in history if e["test_name"] == fc]
            st.markdown(f"<p style='color:rgba(185,222,255,0.85);font-weight:600;margin-bottom:8px;'>{len(filtered)} record(s)</p>", unsafe_allow_html=True)

            for entry in filtered:
                tn, ts, summ = entry["test_name"], entry["timestamp"], entry.get("summary", {})
                pills = []
                for k, v in summ.items():
                    if isinstance(v, (int, float)): pills.append(f"{k}: {round(v, 3)}")
                    elif isinstance(v, str) and len(v) < 80 and k not in ("procedure", "formulas"): pills.append(f"{k}: {v}")
                ph = "".join(f'<span class="hist-prop">{p}</span>' for p in pills[:8])
                st.markdown(f'<div class="hist-card"><div class="hist-test-name">🧪 {tn}</div><div class="hist-time">🕐 {ts}</div>{ph or "<span style=color:rgba(140,180,230,0.6)>No scalar results.</span>"}</div>', unsafe_allow_html=True)

                recs = get_is_recommendations(tn, summ)
                if recs and recs[0][2] != "info":
                    with st.expander(f"📋 IS Code Recommendations — {tn}"):
                        for title, body, level in recs:
                            st.markdown(f'<div class="rec-card {level}"><div class="rec-title">{title}</div><div class="rec-body">{body}</div></div>', unsafe_allow_html=True)

                df_entries = {k: v for k, v in summ.items() if isinstance(v, list)}
                if df_entries:
                    with st.expander(f"📊 Data Table — {tn}"):
                        for label, records in df_entries.items():
                            try:
                                st.markdown(f"**{label}**")
                                st.dataframe(pd.DataFrame(records), width='stretch')
                            except Exception:
                                pass

                with st.expander(f"📤 Share — {tn}"):
                    share_buttons(tn, summ, inside_expander=True)

    # ── TEST VIEW ──
    else:
        result = None
        if selected_module is not None:
            result = selected_module.run()
        else:
            st.info("Test module not available in this environment.")

        if result is not None:
            st.session_state.completed_tests[selected_test] = result
            st.session_state.last_result    = result
            st.session_state.last_test_name = selected_test
            save_history(st.session_state.user_email, selected_test, result)
            st.toast(f"✅ '{selected_test}' saved to history!", icon="🕒")

        if st.session_state.last_result and st.session_state.last_test_name == selected_test:
            res  = st.session_state.last_result
            recs = get_is_recommendations(selected_test, res)
            st.markdown("---")
            st.markdown('<div class="rec-section-title">📋 IS Code Recommendations</div>', unsafe_allow_html=True)
            for title, body, level in recs:
                st.markdown(f'<div class="rec-card {level}"><div class="rec-title">{title}</div><div class="rec-body">{body}</div></div>', unsafe_allow_html=True)

            sym, name_cls, emoji = get_soil_classification(res)
            if sym:
                st.markdown('<div class="rec-section-title" style="margin-top:18px;">🤖 AI Soil Classification (IS 1498 / USCS)</div>', unsafe_allow_html=True)
                st.markdown(f"""
                <div class="class-section">
                    <div class="soil-badge">
                        <span class="soil-badge-symbol">{sym}</span>
                        <span class="soil-badge-name">{emoji} {name_cls}</span>
                    </div>
                    <p style="font-size:0.81rem;color:rgba(185,222,255,0.8);margin-top:7px;">
                        Classification per IS 1498 (USCS method).
                    </p>
                </div>""", unsafe_allow_html=True)

            st.markdown('<div class="share-section">', unsafe_allow_html=True)
            share_buttons(selected_test, res, doc_bytes=build_single_test_docx(selected_test, res))
            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("""
    <button onclick="window.scrollTo({top:0,behavior:'smooth'})"
        style="position:fixed;bottom:24px;right:18px;z-index:9999;
        background:linear-gradient(135deg,rgba(0,100,255,0.92),rgba(0,60,200,1));
        color:#fff;border:2px solid rgba(0,200,255,0.5);border-radius:50%;
        width:48px;height:48px;font-size:1.3rem;font-weight:900;
        box-shadow:0 4px 20px rgba(0,100,255,0.5);cursor:pointer;
        display:flex;align-items:center;justify-content:center;">↑</button>
    """, unsafe_allow_html=True)