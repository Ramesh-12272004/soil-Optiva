import streamlit as st
import numpy as np
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the in-situ dry density of soil using the core cutter method
  as per IS 2720 (Part 29) – 1975, reaffirmed 1995.

Apparatus Required:
  - Cylindrical core cutter (internal diameter ≈ 10 cm, height ≈ 12.5–13 cm)
  - Steel dolley (25 mm high, same diameter as cutter)
  - Steel rammer (≈ 9 kg)
  - Palette knife / trowel
  - Weighing balance (accuracy 1 g)
  - Moisture content cans
  - Oven (105 °C – 110 °C)
  - Straight edge / scraper

Theory:
  The core cutter is driven vertically into the soil, thus extracting an
  undisturbed sample of known volume. The bulk density is calculated from
  the mass of extracted soil and the known internal volume of the cutter.
  The dry density is obtained by dividing the bulk density by (1 + w/100).

  Interpretation (IS 2720):
    ρd < 1.40 g/cc  → Poorly compacted / loose fill
    1.40–1.75 g/cc  → Moderate compaction
    > 1.75 g/cc     → Well compacted

Step-by-Step Procedure:
  1. Clean, dry, and weigh the core cutter (W_empty).
  2. Measure the internal height (H) and diameter (D) of the cutter precisely.
  3. Level the test area; remove any loose surface material.
  4. Place the dolley on top of the cutter and drive both into the ground
     with the rammer until 1–2 cm of dolley remains above ground.
  5. Dig around the cutter with a trowel; carefully lift it out,
     keeping both ends of the soil core intact.
  6. Trim flush with a straight edge; weigh the full cutter (W_full).
  7. Extract a representative soil sample from the core into a moisture can.
     Record W_container, W_wet, and W_dry after oven drying.
  8. Repeat at 3–5 locations for a representative result.

Precautions:
  - Drive the cutter vertically to avoid disturbance.
  - Trim both ends cleanly to ensure the volume equals πD²H/4.
  - Seal the moisture can immediately after extraction.
  - Weigh the wet sample before any drying begins.
"""

FORMULAS = """
Volume of Core Cutter:
  V = (π / 4) × D² × H    [cm³]

Bulk Density:
  ρ_bulk = (W_full – W_empty) / V    [g/cm³]

Moisture Content:
  w (%) = [(W_wet – W_dry) / (W_dry – W_container)] × 100

Dry Density:
  ρ_dry = ρ_bulk / (1 + w / 100)    [g/cm³]

Where:
  D          = Internal diameter of cutter (cm)
  H          = Internal height of cutter (cm)
  W_empty    = Mass of empty cutter (g)
  W_full     = Mass of cutter + compacted soil (g)
  W_container= Mass of empty moisture can (g)
  W_wet      = Mass of can + wet soil (g)
  W_dry      = Mass of can + oven-dried soil (g)
"""


def _generate_report(res, procedure, formulas):
    doc = Document()
    doc.add_heading("Core Cutter In-Situ Density Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 29) – 1975")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the in-situ bulk and dry density of soil using the "
        "core cutter method."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Input Data", 1)
    input_table = doc.add_table(rows=1, cols=2)
    input_table.style = "Table Grid"
    input_table.rows[0].cells[0].text = "Parameter"
    input_table.rows[0].cells[1].text = "Value"
    items = [
        ("Diameter D (cm)",              f"{res['d']:.2f}"),
        ("Height H (cm)",                f"{res['h']:.2f}"),
        ("Volume V (cm³)",               f"{res['volume']:.2f}"),
        ("W_empty – Empty Cutter (g)",   f"{res['w_empty']:.3f}"),
        ("W_full – Cutter + Soil (g)",   f"{res['w_full']:.3f}"),
        ("W_container – Empty Can (g)",  f"{res['w_container']:.3f}"),
        ("W_wet – Can + Wet Soil (g)",   f"{res['w_wet']:.3f}"),
        ("W_dry – Can + Dry Soil (g)",   f"{res['w_dry']:.3f}"),
    ]
    for name, val in items:
        r = input_table.add_row().cells
        r[0].text = name
        r[1].text = val

    doc.add_heading("Results", 1)
    result_table = doc.add_table(rows=1, cols=2)
    result_table.style = "Table Grid"
    result_table.rows[0].cells[0].text = "Parameter"
    result_table.rows[0].cells[1].text = "Value"
    results = [
        ("Bulk Density ρ_bulk (g/cm³)", f"{res['bulk_density']:.3f}"),
        ("Moisture Content w (%)",       f"{res['moisture_content']:.2f}"),
        ("Dry Density ρ_dry (g/cm³)",    f"{res['dry_density']:.3f}"),
        ("Compaction Assessment",        res["suitability"]),
    ]
    for name, val in results:
        r = result_table.add_row().cells
        r[0].text = name
        r[1].text = val

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The in-situ dry density of the soil is {res['dry_density']:.3f} g/cm³ "
        f"and the moisture content is {res['moisture_content']:.2f}%. "
        f"Assessment: {res['suitability']}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🔩 Core Cutter In-Situ Density Test (IS 2720 Part 29 : 1975)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "cc_res" not in st.session_state:
        st.session_state.cc_res = None

    # ── Dimensions ──
    st.markdown("### 📏 Cutter Dimensions")
    c1, c2 = st.columns(2)
    h = c1.number_input("Height H (cm)", value=12.80, min_value=0.1, format="%.2f", key="cc_h")
    d = c2.number_input("Diameter D (cm)", value=10.00, min_value=0.1, format="%.2f", key="cc_d")
    vol = np.pi / 4 * d ** 2 * h
    st.info(f"📐 Calculated Volume V = **{vol:.2f} cm³**")

    # ── Weights ──
    st.markdown("### ⚖️ Weight Measurements")
    col1, col2 = st.columns(2)
    with col1:
        w_empty     = st.number_input("W_empty – Empty Cutter (g)",   value=0.0, min_value=0.0, format="%.3f", key="cc_we")
        w_full      = st.number_input("W_full – Cutter + Soil (g)",   value=0.0, min_value=0.0, format="%.3f", key="cc_wf")
    with col2:
        w_container = st.number_input("W_container – Empty Can (g)",  value=0.0, min_value=0.0, format="%.3f", key="cc_wc")
        w_wet       = st.number_input("W_wet – Can + Wet Soil (g)",   value=0.0, min_value=0.0, format="%.3f", key="cc_ww")
        w_dry       = st.number_input("W_dry – Can + Dry Soil (g)",   value=0.0, min_value=0.0, format="%.3f", key="cc_wd")

    if st.button("📊 Calculate"):
        errors = []
        if w_full <= w_empty:
            errors.append("W_full must be greater than W_empty.")
        if w_wet <= w_container:
            errors.append("W_wet must be greater than W_container.")
        if w_dry <= w_container:
            errors.append("W_dry must be greater than W_container.")
        if w_dry >= w_wet:
            errors.append("W_dry must be less than W_wet.")

        if errors:
            for e in errors:
                st.error(e)
            return None

        bulk_density     = (w_full - w_empty) / vol
        Wd               = w_dry - w_container
        Ww               = w_wet - w_container
        moisture_content = (Ww - Wd) / Wd * 100
        dry_density      = bulk_density / (1 + moisture_content / 100)

        if dry_density < 1.40:
            suitability = "Poorly compacted / loose fill — compaction required."
        elif dry_density < 1.75:
            suitability = "Moderate compaction — acceptable for general earthwork."
        else:
            suitability = "Well compacted — suitable for structural fill."

        st.session_state.cc_res = {
            "h": h, "d": d, "volume": vol,
            "w_empty": w_empty, "w_full": w_full,
            "w_container": w_container, "w_wet": w_wet, "w_dry": w_dry,
            "bulk_density": bulk_density,
            "moisture_content": moisture_content,
            "dry_density": dry_density,
            "suitability": suitability,
        }

    if st.session_state.cc_res is not None:
        res = st.session_state.cc_res
        st.markdown("### 📋 Results")
        r1, r2, r3 = st.columns(3)
        r1.metric("Bulk Density (g/cm³)",    f"{res['bulk_density']:.3f}")
        r2.metric("Moisture Content (%)",    f"{res['moisture_content']:.2f}")
        r3.metric("Dry Density (g/cm³)",     f"{res['dry_density']:.3f}")
        st.info(f"🏷️ **Assessment:** {res['suitability']}")

        report_buf = _generate_report(res, PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Core_Cutter_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":                PROCEDURE,
            "formulas":                 FORMULAS,
            "Bulk Density (g/cm3)":     round(res["bulk_density"], 3),
            "Moisture Content (%)":     round(res["moisture_content"], 2),
            "Dry Density (g/cm3)":      round(res["dry_density"], 3),
            "Compaction Assessment":    res["suitability"],
        }

    return None