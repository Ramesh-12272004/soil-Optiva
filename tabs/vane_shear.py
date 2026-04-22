import streamlit as st
import math
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the undrained shear strength (Su) and sensitivity of soft
  cohesive soils using the vane shear test as per IS 2720 (Part 30) – 1980.

Apparatus Required:
  - Vane shear apparatus (field or lab type)
  - Four-bladed vane (rectangular: D ≈ 12 mm, H ≈ 25 mm for lab; larger for field)
  - Torque measuring device (spring balance or torque wrench) with graduated scale
  - Stopwatch

Theory:
  A four-bladed vane is pushed into the soil and rotated at a slow, constant
  rate (≈ 0.1°/s). The torque T required to shear the cylindrical surface
  (top, bottom, and perimeter of the vane) provides the undrained shear strength.

  If both undisturbed and remoulded strengths are measured, the ratio gives
  the Sensitivity:
    S_t = S_u(undisturbed) / S_u(remoulded)

  Sensitivity Classification (IS):
    S_t < 2      → Insensitive
    2 – 4        → Normal sensitive
    4 – 8        → Sensitive
    8 – 16       → Extra sensitive
    > 16         → Quick clay

Step-by-Step Procedure:
  1. Measure the vane diameter D and height H precisely.
  2. Determine the spring constant of the torquemeter device.
  3. Insert the vane into the undisturbed soil at the required depth without
     pre-shearing; apply light vertical pressure only.
  4. Rotate the vane at ≈ 0.1°/s (about 6°/min); record the angle of twist
     at failure (maximum torque).
  5. Continue rotating rapidly (≥ 5 turns) to fully remould the soil.
  6. Re-measure at the same rate to obtain remoulded torque.
  7. Repeat at a minimum of 2–3 depths or locations.

Precautions:
  - Insert the vane without rotation to minimise disturbance.
  - Begin the test within 5 minutes of insertion.
  - Maintain a constant angular rotation rate.
  - Avoid surcharging the test area before testing.
"""

FORMULAS = """
Vane Constant (K_v):
  K_v = π × D² × H × (1/2 + D/(6H))    [cm³]

Torque:
  T = Spring Constant × (Final Angle – Initial Angle)    [kg·cm]

Shear Strength:
  S_u = T / K_v    [kg/cm²]

Sensitivity:
  S_t = S_u (Undisturbed) / S_u (Remoulded)

Where:
  D = Diameter of vane (cm)
  H = Height of vane (cm)
  Spring Constant in kg·cm/degree
"""


def _sensitivity_class(st_val):
    if st_val < 2:
        return "Insensitive"
    elif st_val < 4:
        return "Normal Sensitive"
    elif st_val < 8:
        return "Sensitive"
    elif st_val < 16:
        return "Extra Sensitive"
    else:
        return "Quick Clay"


def _generate_report(D, H, Kv, spring_const, df, avg_su, sensitivity_val, sens_class, procedure, formulas):
    doc = Document()
    doc.add_heading("Vane Shear Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 30) – 1980")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the undrained shear strength and sensitivity of soft "
        "cohesive soil using the vane shear test."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Vane Diameter D = {D:.2f} cm")
    doc.add_paragraph(f"Vane Height   H = {H:.2f} cm")
    doc.add_paragraph(f"Vane Constant Kv = {Kv:.4f} cm³")
    doc.add_paragraph(f"Spring Constant = {spring_const:.4f} kg·cm/degree")

    doc.add_heading("Results Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.4f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Average Shear Strength Su = {avg_su:.4f} kg/cm²")
    if sensitivity_val != "N/A":
        doc.add_paragraph(f"Sensitivity St = {sensitivity_val:.2f}  ({sens_class})")

    doc.add_heading("Conclusion", 1)
    conc = (
        f"The undrained shear strength of the soil is {avg_su:.4f} kg/cm²."
    )
    if sensitivity_val != "N/A":
        conc += f" The sensitivity is {sensitivity_val:.2f} ({sens_class})."
    doc.add_paragraph(conc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🌀 Vane Shear Test (IS 2720 Part 30 : 1980)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "vs_res" not in st.session_state:
        st.session_state.vs_res = None

    # ── Apparatus ──
    st.markdown("### 📏 Vane Dimensions & Spring Constant")
    c1, c2, c3 = st.columns(3)
    D   = c1.number_input("Diameter D (cm)", value=1.20, min_value=0.01, format="%.3f", key="vs_D")
    H   = c2.number_input("Height H (cm)",   value=2.40, min_value=0.01, format="%.3f", key="vs_H")
    sc  = c3.number_input("Spring Constant (kg·cm/degree)", value=0.001, min_value=0.0001, format="%.4f", key="vs_sc")

    Kv = math.pi * D ** 2 * H * (0.5 + D / (6 * H))
    st.info(f"Vane Constant K_v = **{Kv:.4f} cm³**")

    # ── Trial inputs ──
    num_trials = st.number_input("Number of Trials", min_value=1, max_value=6, value=2, step=1, key="vs_nt")

    if "vs_inputs" not in st.session_state or len(st.session_state.vs_inputs) != num_trials:
        st.session_state.vs_inputs = [
            {"type": "Undisturbed", "init_deg": 0.0, "fail_deg": 0.0}
            for _ in range(num_trials)
        ]

    st.markdown("### 📋 Angle of Twist Readings")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3 = st.columns(3)
        inp = st.session_state.vs_inputs[i]
        inp["type"] = c1.selectbox(
            "Sample Type", ["Undisturbed", "Remoulded"],
            index=0 if inp["type"] == "Undisturbed" else 1,
            key=f"vs_type_{i}"
        )
        inp["init_deg"] = c2.number_input(
            "Initial Reading (°)", value=inp["init_deg"],
            min_value=0.0, format="%.2f", key=f"vs_init_{i}"
        )
        inp["fail_deg"] = c3.number_input(
            "Failure Reading (°)", value=inp["fail_deg"],
            min_value=0.0, format="%.2f", key=f"vs_fail_{i}"
        )

    if st.button("🔄 Reset"):
        st.session_state.vs_inputs = [
            {"type": "Undisturbed", "init_deg": 0.0, "fail_deg": 0.0}
            for _ in range(num_trials)
        ]
        st.session_state.vs_res = None
        st.rerun()

    if st.button("📊 Calculate Shear Strength"):
        rows = []
        su_undisturbed = None
        su_remoulded   = None

        for i, inp in enumerate(st.session_state.vs_inputs[:num_trials]):
            diff = inp["fail_deg"] - inp["init_deg"]
            if diff <= 0:
                st.warning(f"Trial {i + 1}: Failure reading must be greater than initial reading.")
                continue
            T  = sc * diff
            Su = T / Kv if Kv > 0 else 0.0
            rows.append({
                "Trial":               i + 1,
                "Type":                inp["type"],
                "Initial (°)":         round(inp["init_deg"], 2),
                "Failure (°)":         round(inp["fail_deg"], 2),
                "Δθ (°)":              round(diff, 2),
                "Torque T (kg·cm)":    round(T, 4),
                "Su (kg/cm²)":         round(Su, 5),
            })
            if inp["type"] == "Undisturbed":
                su_undisturbed = Su
            elif inp["type"] == "Remoulded":
                su_remoulded = Su

        if not rows:
            st.error("No valid trials calculated.")
            return None

        df = pd.DataFrame(rows)
        avg_su = float(df["Su (kg/cm²)"].mean())

        sensitivity_val = "N/A"
        sens_class      = "N/A"
        if su_undisturbed and su_remoulded and su_remoulded > 0:
            sensitivity_val = su_undisturbed / su_remoulded
            sens_class      = _sensitivity_class(sensitivity_val)

        st.markdown("### 📋 Results Table")
        st.dataframe(df, use_container_width=True)

        c1, c2 = st.columns(2)
        c1.metric("Average Su (kg/cm²)", f"{avg_su:.5f}")
        if sensitivity_val != "N/A":
            c2.metric("Sensitivity St", f"{sensitivity_val:.2f}")
            st.info(f"🏷️ Sensitivity Class: **{sens_class}**")

        st.session_state.vs_res = {
            "df": df, "avg_su": avg_su,
            "sensitivity_val": sensitivity_val,
            "sens_class": sens_class,
        }

        report_buf = _generate_report(
            D, H, Kv, sc, df, avg_su,
            sensitivity_val, sens_class,
            PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Vane_Shear_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        result = {
            "procedure":            PROCEDURE,
            "formulas":             FORMULAS,
            "data":                 df,
            "Average Su (kg/cm2)":  round(avg_su, 5),
        }
        if sensitivity_val != "N/A":
            result["Sensitivity St"]      = round(sensitivity_val, 2)
            result["Sensitivity Class"]   = sens_class
        return result

    return None