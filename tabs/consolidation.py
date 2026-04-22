import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the consolidation characteristics of fine-grained soil –
  specifically the Compression Index (Cc), Coefficient of Volume
  Compressibility (mv), and Coefficient of Consolidation (cv) –
  as per IS 2720 (Part 15) – 1986.

Apparatus Required:
  - Fixed-ring consolidation cell (oedometer)
  - Loading device with hanger and dead weights
  - Dial gauge (0.01 mm / div) or LVDT
  - Stopwatch
  - Balance (accuracy 0.1 g)
  - Sample trimmer and cutting ring
  - Porous stones, filter paper

Theory:
  When a saturated soil is loaded, excess pore water pressure develops.
  As water drains out, the soil compresses – this is called consolidation.
  The test is performed by applying load increments (each double the previous)
  and recording time–settlement for each increment.

  Key parameters:
    Cc  – slope of the e–log(P) curve (compression index); indicates
          compressibility under increasing load.
    Cs  – swelling/recompression index (slope on unloading).
    mv  – coefficient of volume compressibility (cm²/kg).
    cv  – coefficient of consolidation (cm²/s or cm²/year).
    
  Classification (Cc values):
    Cc < 0.1  → Low compressibility
    0.1–0.3   → Moderate compressibility
    > 0.3     → High compressibility (organic / soft clay)

Step-by-Step Procedure:
  1. Prepare a thin, undisturbed specimen fitting the consolidation ring
     (typical H ≈ 20 mm, D ≈ 60–75 mm).
  2. Record initial height H₀ and diameter D; calculate area A.
  3. Estimate initial void ratio from bulk/dry density and Gs.
  4. Place specimen between wet porous stones; flood the cell.
  5. Apply an initial seating load (≈ 5 kPa); record initial dial reading.
  6. Apply load increments (e.g., 10, 20, 40, 80, 160, 320 kPa);
     record final dial gauge reading for each after full consolidation
     (≥ 24 h per increment for clays).
  7. For each increment, also record time–settlement data (log-t or √t
     method) to determine cv.
  8. On completion, record final reading; back-calculate final void ratio.

Precautions:
  - Do not allow the specimen to dry during the test.
  - Keep the load arm balanced; apply loads gently without impact.
  - Ensure full saturation; maintain water level in the cell at all times.
  - Trim the specimen carefully to avoid disturbance.
"""

FORMULAS = """
Cross-sectional Area:
  A = (π / 4) × D²    [cm²]

Compression (Settlement) per increment:
  ΔH = (Final Dial Reading – Initial Dial Reading) × Least Count    [mm]

Axial Strain:
  ε = ΔH / H₀

Change in Void Ratio:
  Δe = ε × (1 + e₀)

Current Void Ratio:
  e = e₀ – Δe    (cumulative from start)

Compression Index (Cc):
  Cc = –Δe / Δlog(P)  [slope of the straight-line portion of e–log(P) curve]

Coefficient of Volume Compressibility (mv):
  mv = Δε / Δσ = Δe / [(1 + e₀) × Δσ]    [cm²/kg]

Where:
  H₀    = Initial specimen height (cm)
  e₀    = Initial void ratio
  D     = Diameter of specimen (cm)
  Δσ    = Stress increment (kg/cm²)
"""


def _generate_report(h0, d, dial_lc, e0, df, cc, mv, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("Consolidation Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 15) – 1986")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the consolidation characteristics (Cc, mv) "
        "of fine-grained soil."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Specimen Height H₀ = {h0:.3f} cm")
    doc.add_paragraph(f"Specimen Diameter D = {d:.3f} cm")
    doc.add_paragraph(f"Dial Gauge LC = {dial_lc:.4f} mm/div")
    doc.add_paragraph(f"Initial Void Ratio e₀ = {e0:.3f}")

    doc.add_heading("Consolidation Data Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.5f}" if isinstance(val, float) else str(val)

    doc.add_heading("Results", 1)
    if cc is not None:
        doc.add_paragraph(f"Compression Index Cc = {cc:.4f}")
    if mv is not None:
        doc.add_paragraph(f"Coefficient of Volume Compressibility mv = {mv:.6f} cm²/kg")

    doc.add_heading("e–log(P) Curve", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    if cc:
        if cc < 0.1:
            comp = "low compressibility"
        elif cc < 0.3:
            comp = "moderate compressibility"
        else:
            comp = "high compressibility"
        doc.add_paragraph(
            f"Compression Index Cc = {cc:.4f}, indicating {comp}. "
            f"mv = {mv:.6f} cm²/kg."
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("📉 Consolidation Test (IS 2720 Part 15 : 1986)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "consol_res" not in st.session_state:
        st.session_state.consol_res = None

    # ── Specimen details ──
    st.markdown("### 📏 Specimen Details")
    c1, c2, c3, c4 = st.columns(4)
    h0      = c1.number_input("H₀ – Initial Height (cm)", value=2.0, min_value=0.01, format="%.3f", key="cs_h0")
    d       = c2.number_input("D – Diameter (cm)",         value=6.0, min_value=0.01, format="%.3f", key="cs_d")
    dial_lc = c3.number_input("Dial Gauge LC (mm/div)",    value=0.01, min_value=0.0001, format="%.4f", key="cs_dlc")
    e0      = c4.number_input("Initial Void Ratio e₀",     value=0.80, min_value=0.0, format="%.3f", key="cs_e0")
    A = np.pi / 4 * d ** 2
    st.info(f"Cross-sectional Area A = **{A:.3f} cm²**")

    # ── Load increments ──
    num_inc = st.number_input("Number of Load Increments", min_value=2, max_value=20, value=6, step=1, key="cs_ni")

    if "cs_inputs" not in st.session_state or len(st.session_state.cs_inputs) != num_inc:
        st.session_state.cs_inputs = [
            {"load": 0.0, "init_div": 0.0, "final_div": 0.0}
            for _ in range(num_inc)
        ]

    st.markdown("### 📋 Load Increment Data")
    for i in range(num_inc):
        c1, c2, c3 = st.columns(3)
        inp = st.session_state.cs_inputs[i]
        inp["load"]      = c1.number_input(f"Load (kg/cm²) [{i+1}]",          value=inp["load"],      min_value=0.0, format="%.4f", key=f"cs_ld_{i}")
        inp["init_div"]  = c2.number_input(f"Initial Dial Reading (div) [{i+1}]", value=inp["init_div"], min_value=0.0, format="%.2f", key=f"cs_id_{i}")
        inp["final_div"] = c3.number_input(f"Final Dial Reading (div) [{i+1}]",   value=inp["final_div"], min_value=0.0, format="%.2f", key=f"cs_fd_{i}")

    if st.button("🔄 Reset"):
        st.session_state.cs_inputs = [{"load": 0.0, "init_div": 0.0, "final_div": 0.0} for _ in range(num_inc)]
        st.session_state.consol_res = None
        st.rerun()

    if st.button("📊 Calculate Consolidation Results"):
        rows = []
        cumulative_e = e0
        for i, inp in enumerate(st.session_state.cs_inputs[:num_inc]):
            load      = inp["load"]
            init_div  = inp["init_div"]
            final_div = inp["final_div"]

            settlement_mm = (final_div - init_div) * dial_lc
            compression   = settlement_mm / 10          # mm → cm
            strain        = compression / h0 if h0 > 0 else 0.0
            delta_e       = strain * (1 + e0)
            cumulative_e -= delta_e
            log_load      = np.log10(load) if load > 0 else np.nan

            rows.append({
                "Load (kg/cm²)":        load,
                "Init Div":             init_div,
                "Final Div":            final_div,
                "Settlement (mm)":      round(settlement_mm, 4),
                "Compression (cm)":     round(compression, 5),
                "Strain ε":             round(strain, 5),
                "Δe":                   round(delta_e, 5),
                "Void Ratio e":         round(cumulative_e, 5),
                "log(P)":               round(log_load, 5) if not np.isnan(log_load) else np.nan,
            })

        df = pd.DataFrame(rows)
        df_valid = df[df["log(P)"].notna() & (df["Load (kg/cm²)"] > 0)].copy()

        if len(df_valid) < 2:
            st.error("At least 2 valid load increments (> 0) needed.")
            return None

        # Cc from last 3 points (virgin compression)
        cc, mv = None, None
        if len(df_valid) >= 3:
            last = df_valid.tail(3)
            slope, _ = np.polyfit(last["log(P)"], last["Void Ratio e"], 1)
            cc = -slope   # negative slope (e decreases with load)

        # mv from first two valid increments
        if len(df_valid) >= 2:
            de     = df_valid["Δe"].iloc[1]
            dsigma = df_valid["Load (kg/cm²)"].iloc[1] - df_valid["Load (kg/cm²)"].iloc[0]
            mv = de / ((1 + e0) * dsigma) if dsigma != 0 else None

        # Plot e–log(P)
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(df_valid["log(P)"], df_valid["Void Ratio e"],
                marker="o", color="#0a68cc", linewidth=2)
        ax.set_xlabel("log(P)  [P in kg/cm²]", fontsize=11)
        ax.set_ylabel("Void Ratio e", fontsize=11)
        ax.set_title("e – log(P) Consolidation Curve", fontsize=13, fontweight="bold")
        ax.invert_yaxis()
        ax.grid(True, linestyle="--", alpha=0.5)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.markdown("### 📋 Consolidation Data Table")
        st.dataframe(df.round(5), use_container_width=True)

        c1, c2 = st.columns(2)
        if cc is not None:
            c1.metric("Compression Index Cc", f"{cc:.4f}")
        if mv is not None:
            c2.metric("mv (cm²/kg)", f"{mv:.6f}")

        if cc is not None:
            if cc < 0.1:
                comp = "Low compressibility – suitable for light structures."
            elif cc < 0.3:
                comp = "Moderate compressibility – design with attention to settlement."
            else:
                comp = "High compressibility – significant settlements expected; consolidation surcharge recommended."
            st.info(f"🏷️ **{comp}**")

        st.image(img_buf)

        st.session_state.consol_res = {
            "df": df, "cc": cc, "mv": mv, "img_buf": img_buf,
            "h0": h0, "d": d, "dial_lc": dial_lc, "e0": e0,
        }

        report_buf = _generate_report(h0, d, dial_lc, e0, df, cc, mv, img_buf, PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Consolidation_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        result = {
            "procedure": PROCEDURE,
            "formulas":  FORMULAS,
            "data":      df,
            "graph":     img_buf,
        }
        if cc is not None:
            result["Compression Index Cc"] = round(cc, 4)
        if mv is not None:
            result["mv (cm2/kg)"] = round(mv, 6)
        return result

    return None