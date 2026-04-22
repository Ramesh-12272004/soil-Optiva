import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Circle
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the shear strength parameters (cohesion c and angle of
  internal friction ϕ) of soil using the undrained triaxial test as per
  IS 2720 (Part 11) – 1971.

Apparatus Required:
  - Triaxial cell with top platen and base pedestal
  - Cell pressure pump / air-water interface cylinder
  - Axial load frame with proving ring and dial gauge
  - Back pressure panel (for UU / CU tests)
  - Rubber membrane and O-rings
  - Weighing balance, moisture cans

Theory:
  In the Unconsolidated Undrained (UU) test, the specimen is subjected to
  a cell (confining) pressure σ₃ without drainage. An axial deviator load
  is then applied until failure. The total principal stresses at failure:
    σ₁ = σ₃ + (σ₁ – σ₃)_failure
  
  Mohr's circles are drawn for each trial. The failure envelope tangent
  to these circles defines c and ϕ.

  For UU tests on saturated clay, ϕ ≈ 0° and c ≈ Su (undrained shear strength).

  Typical shear strength parameters:
    Dense sand:   c = 0, ϕ = 35–45°
    Loose sand:   c = 0, ϕ = 28–34°
    Clay (UU):    c = Su, ϕ ≈ 0°

Step-by-Step Procedure:
  1. Prepare cylindrical specimen (H/D ≈ 2); measure D and L.
  2. Encase in rubber membrane; seat on pedestal; fit top cap.
  3. Assemble the triaxial cell; fill with water; apply σ₃.
  4. Apply the axial deviator stress at constant strain rate (1% / min).
  5. Record proving ring reading at each 0.5% axial strain increment.
  6. Stop at failure (peak deviator stress) or 20% axial strain.
  7. Repeat for at least 3 different σ₃ values.
  8. Draw Mohr's circles; fit the failure envelope.

Precautions:
  - Ensure no air remains in the cell or specimen drainage lines.
  - Maintain the strain rate constant throughout.
  - Inspect membrane for punctures before assembly.
  - Record all readings to the nearest division.
"""

FORMULAS = """
Deviator Stress at failure:
  (σ₁ – σ₃) = Failure Load P / Corrected Area A_c

Principal Stresses:
  σ₃ = Cell (confining) pressure    [kg/cm²]
  σ₁ = σ₃ + Deviator Stress         [kg/cm²]

Mohr Circle:
  Centre = (σ₁ + σ₃) / 2
  Radius = (σ₁ – σ₃) / 2

Failure Envelope (linear regression):
  τ = c + σ × tan(ϕ)

  Using centre-radius points:
    Radius ≈ c × cos(ϕ) + Centre × sin(ϕ)   (exact tangent condition)
  
  Approximate method (tangent to circles at top):
    Fit linear regression through (Centre, Radius) pairs:
      slope → sin(ϕ)    ⟹  ϕ = arcsin(slope)
      intercept → c × cos(ϕ)   ⟹  c = intercept / cos(ϕ)

Undrained Shear Strength (for ϕ ≈ 0):
  Su = (σ₁ – σ₃) / 2
"""


def _generate_report(df, cohesion, phi, mean_su, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("Undrained Triaxial Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 11) – 1971")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine cohesion c and angle of internal friction ϕ "
        "from Mohr's circles and the failure envelope."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Trial Data Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.4f}" if isinstance(val, float) else str(val)

    doc.add_heading("Results", 1)
    doc.add_paragraph(f"Cohesion c = {cohesion:.4f} kg/cm²")
    doc.add_paragraph(f"Angle of Internal Friction ϕ = {phi:.2f}°")
    doc.add_paragraph(f"Average Su (UU) = {mean_su:.4f} kg/cm²")

    doc.add_heading("Mohr's Circles", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The soil has cohesion c = {cohesion:.4f} kg/cm² and "
        f"angle of internal friction ϕ = {phi:.2f}°. "
        f"Average undrained shear strength Su = {mean_su:.4f} kg/cm²."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("⭕ Undrained Triaxial Test (IS 2720 Part 11 : 1971)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "tri_res" not in st.session_state:
        st.session_state.tri_res = None

    num_trials = st.number_input("Number of Trials", min_value=2, max_value=8, value=3, step=1, key="tri_nt")

    if "tri_inputs" not in st.session_state or len(st.session_state.tri_inputs) != num_trials:
        st.session_state.tri_inputs = [
            {"sigma3": 0.0, "dev": 0.0}
            for _ in range(num_trials)
        ]

    st.markdown("### 📋 Enter Stress Data")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2 = st.columns(2)
        inp = st.session_state.tri_inputs[i]
        inp["sigma3"] = c1.number_input(
            f"Cell Pressure σ₃ (kg/cm²)", value=inp["sigma3"],
            min_value=0.0, format="%.3f", key=f"tri_s3_{i}"
        )
        inp["dev"] = c2.number_input(
            f"Deviator Stress (σ₁–σ₃) at failure (kg/cm²)", value=inp["dev"],
            min_value=0.0, format="%.3f", key=f"tri_dev_{i}"
        )

    if st.button("🔄 Reset"):
        st.session_state.tri_inputs = [{"sigma3": 0.0, "dev": 0.0} for _ in range(num_trials)]
        st.session_state.tri_res = None
        st.rerun()

    if st.button("📊 Calculate & Draw Mohr's Circles"):
        rows = []
        for i, inp in enumerate(st.session_state.tri_inputs[:num_trials]):
            s3  = inp["sigma3"]
            dev = inp["dev"]
            s1  = s3 + dev
            c   = (s1 + s3) / 2
            r   = (s1 - s3) / 2
            su  = r
            rows.append({
                "Trial": i + 1,
                "σ₃ (kg/cm²)": s3,
                "Deviator (σ₁–σ₃) (kg/cm²)": dev,
                "σ₁ (kg/cm²)": round(s1, 4),
                "Centre (kg/cm²)": round(c, 4),
                "Radius (kg/cm²)": round(r, 4),
                "Su (kg/cm²)": round(su, 4),
            })

        df = pd.DataFrame(rows)
        centres = df["Centre (kg/cm²)"].values
        radii   = df["Radius (kg/cm²)"].values
        mean_su = float(df["Su (kg/cm²)"].mean())

        # Fit failure envelope: R = a + b*C  → sin(ϕ)=b, c*cos(ϕ)=a
        coeffs  = np.polyfit(centres, radii, 1)
        b, a    = coeffs  # b = slope, a = intercept
        sin_phi = b
        phi_rad = np.arcsin(np.clip(sin_phi, -1, 1))
        phi_deg = np.degrees(phi_rad)
        cohesion = a / np.cos(phi_rad) if np.cos(phi_rad) != 0 else a

        st.success(f"**Cohesion c = {cohesion:.4f} kg/cm²**")
        st.success(f"**Angle of Internal Friction ϕ = {phi_deg:.2f}°**")
        st.info(f"**Average Su (UU) = {mean_su:.4f} kg/cm²**")

        # ── Mohr's circle plot ──
        fig, ax = plt.subplots(figsize=(9, 5))
        all_sigma1 = df["σ₁ (kg/cm²)"].values
        all_sigma3 = df["σ₃ (kg/cm²)"].values
        x_max = float(max(all_sigma1)) * 1.15
        ax.set_xlim(0, max(x_max, 5))
        ax.set_ylim(0, max(radii) * 1.6)

        for c_val, r_val in zip(centres, radii):
            circle = Circle((c_val, 0), r_val, fill=False, edgecolor="#0a68cc",
                            linewidth=2)
            ax.add_patch(circle)

        # Failure envelope
        x_env = np.linspace(0, x_max, 200)
        y_env = cohesion + x_env * np.tan(phi_rad)
        ax.plot(x_env, y_env, color="#e05c00", linewidth=2,
                label=f"τ = {cohesion:.3f} + σ·tan({phi_deg:.1f}°)")
        ax.set_aspect("equal", adjustable="datalim")
        ax.set_xlabel("Normal Stress σ (kg/cm²)")
        ax.set_ylabel("Shear Stress τ (kg/cm²)")
        ax.set_title("Mohr's Circles – Triaxial Test")
        ax.legend()
        ax.grid(True, linestyle="--", alpha=0.5)
        st.pyplot(fig)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.markdown("### 📋 Trial Data")
        st.dataframe(df.round(4), use_container_width=True)

        st.session_state.tri_res = {
            "df": df, "cohesion": cohesion, "phi": phi_deg,
            "mean_su": mean_su, "img_buf": img_buf
        }

        report_buf = _generate_report(df, cohesion, phi_deg, mean_su, img_buf, PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Triaxial_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":                    PROCEDURE,
            "formulas":                     FORMULAS,
            "data":                         df,
            "graph":                        img_buf,
            "Cohesion c (kg/cm2)":          round(cohesion, 4),
            "Friction Angle phi (degrees)": round(phi_deg, 2),
            "Average Su (kg/cm2)":          round(mean_su, 4),
        }

    return None