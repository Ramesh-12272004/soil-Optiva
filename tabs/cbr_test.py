import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the California Bearing Ratio (CBR) of compacted soil for
  pavement design as per IS 2720 (Part 16) – 1987.

Apparatus Required:
  - CBR mould (150 mm internal diameter, 175 mm height, volume ≈ 2250 cm³)
  - Spacer disc (148 mm diameter, 47.7 mm high)
  - Standard rammer and collar
  - Perforated base plate and surcharge weights (each 2.5 kg)
  - Penetration piston (50 mm diameter)
  - Load frame with proving ring and dial gauge
  - Soaking tank (for soaked CBR)
  - Stopwatch

Theory:
  CBR is the ratio of the load required to penetrate the soil at a specified
  rate to the load required to penetrate a standard crushed stone at the same
  rate. It is expressed as a percentage.

  Standard loads (from IS 2720 Part 16):
    At 2.5 mm penetration → 1370 kg
    At 5.0 mm penetration → 2055 kg

  If the value at 5.0 mm > value at 2.5 mm, repeat the test.
  If the same result is obtained on re-testing, use the 5.0 mm value as CBR.

  CBR Usage in Pavement Design:
    CBR < 3   → Very weak subgrade; heavy stabilisation required
    3–7       → Poor subgrade; thick pavement layers needed
    7–15      → Moderate subgrade; suitable for light traffic
    15–30     → Good subgrade; flexible pavement
    > 30      → Excellent subgrade

Step-by-Step Procedure:
  1. Compact the soil specimen at OMC in the CBR mould (3 layers, 56 blows each).
  2. Remove collar; trim flush; fit base plate and surcharge weights.
  3. For soaked CBR: soak in water for 96 h (4 days) with 2.5 kg surcharge on top.
  4. Remove from tank; drain for 15 min; then test.
  5. Place the mould on the loading frame; lower the piston onto the soil.
  6. Apply a seating load of 4 kg; zero both gauges.
  7. Apply load at 1.25 mm/min; record load at penetrations of
     0.5, 1.0, 1.5, 2.0, 2.5, 3.0, 4.0, 5.0, 7.5, 10.0, 12.5 mm.
  8. Remove the mould; take moisture content from top and bottom.

Precautions:
  - The surcharge weight during soaking must equal that during testing.
  - Maintain a constant penetration rate of 1.25 mm/min.
  - Check for any unusual stress–penetration shapes (concave upward curve).
  - Apply a correction if the load–penetration curve shows a concave start.
"""

FORMULAS = """
Test Load at each penetration:
  P_test = Dial Reading (div) × Proving Ring Constant (kg/div)    [kg]

CBR at 2.5 mm:
  CBR₂.₅ (%) = (P_test at 2.5 mm / 1370) × 100

CBR at 5.0 mm:
  CBR₅.₀ (%) = (P_test at 5.0 mm / 2055) × 100

Final CBR:
  CBR = max(CBR₂.₅, CBR₅.₀)  [use CBR₅.₀ only if consistently > CBR₂.₅ on repeat tests]

Standard Loads (IS 2720 Part 16):
  Penetration 2.5 mm → 1370 kg (from crushed stone standard)
  Penetration 5.0 mm → 2055 kg
"""


def _generate_report(ring_const, df, cbr_25, cbr_50, final_cbr, conclusion, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("CBR (California Bearing Ratio) Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 16) – 1987")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the CBR value of compacted soil for flexible pavement design."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Proving Ring Constant = {ring_const:.4f} kg/div")

    doc.add_heading("Load–Penetration Data", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.3f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"CBR at 2.5 mm = {cbr_25:.2f} %")
    doc.add_paragraph(f"CBR at 5.0 mm = {cbr_50:.2f} %")
    doc.add_paragraph(f"Final CBR     = {final_cbr:.2f} %")

    doc.add_heading("Load–Penetration Curve", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"Final CBR = {final_cbr:.2f}%. {conclusion}"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🛣️ California Bearing Ratio (CBR) Test (IS 2720 Part 16 : 1987)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "cbr_res" not in st.session_state:
        st.session_state.cbr_res = None

    st.markdown("### 🔧 Proving Ring Constant")
    ring_const = st.number_input("Proving Ring Constant (kg/div)", value=1.0, min_value=0.0001,
                                 format="%.4f", key="cbr_rc")

    # Standard penetrations
    standard_pens = [0.5, 1.0, 1.5, 2.0, 2.5, 3.0, 4.0, 5.0, 7.5, 10.0, 12.5]

    # Initialize input storage
    if "cbr_inputs" not in st.session_state:
        st.session_state.cbr_inputs = {p: 0.0 for p in standard_pens}

    st.markdown("### 📋 Enter Proving Ring Dial Readings")
    st.caption("Enter dial readings (in divisions) at each standard penetration depth.")
    cols_per_row = 4
    pen_list = standard_pens
    rows_needed = (len(pen_list) + cols_per_row - 1) // cols_per_row

    for row_i in range(rows_needed):
        cols = st.columns(cols_per_row)
        for col_i, pen in enumerate(pen_list[row_i * cols_per_row:(row_i + 1) * cols_per_row]):
            st.session_state.cbr_inputs[pen] = cols[col_i].number_input(
                f"Pen {pen} mm (div)", value=st.session_state.cbr_inputs.get(pen, 0.0),
                min_value=0.0, format="%.2f", key=f"cbr_pen_{pen}"
            )

    if st.button("🔄 Reset"):
        st.session_state.cbr_inputs = {p: 0.0 for p in standard_pens}
        st.session_state.cbr_res = None
        st.rerun()

    if st.button("📊 Calculate CBR"):
        rows = []
        for pen in standard_pens:
            dial = st.session_state.cbr_inputs.get(pen, 0.0)
            load = dial * ring_const
            rows.append({
                "Penetration (mm)": pen,
                "Dial Reading (div)": dial,
                "Load (kg)": round(load, 3),
            })

        df = pd.DataFrame(rows)

        # Check required penetrations
        loads_by_pen = {r["Penetration (mm)"]: r["Load (kg)"] for _, r in df.iterrows()}
        if 2.5 not in loads_by_pen or 5.0 not in loads_by_pen:
            st.error("Penetration data must include 2.5 mm and 5.0 mm readings.")
            return None

        load_25 = loads_by_pen[2.5]
        load_50 = loads_by_pen[5.0]

        cbr_25   = (load_25 / 1370) * 100
        cbr_50   = (load_50 / 2055) * 100
        final_cbr = max(cbr_25, cbr_50)

        if final_cbr < 3:
            conclusion = "Very weak subgrade. Heavy stabilisation / geo-grid reinforcement required."
        elif final_cbr < 7:
            conclusion = "Poor subgrade. Thick granular sub-base layers required."
        elif final_cbr < 15:
            conclusion = "Moderate subgrade. Suitable for low-volume roads with proper pavement design."
        elif final_cbr < 30:
            conclusion = "Good subgrade. Suitable for flexible pavement under moderate traffic."
        else:
            conclusion = "Excellent subgrade. Suitable for heavy traffic pavement with thin sections."

        # Plot
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(df["Penetration (mm)"], df["Load (kg)"],
                marker="o", color="#0a68cc", linewidth=2, label="Load–Penetration")
        ax.axvline(2.5, color="#888", linestyle="--", linewidth=1, alpha=0.7, label="2.5 mm")
        ax.axvline(5.0, color="#aaa", linestyle="--", linewidth=1, alpha=0.7, label="5.0 mm")
        ax.set_xlabel("Penetration (mm)")
        ax.set_ylabel("Load (kg)")
        ax.set_title("Load–Penetration Curve")
        ax.legend()
        ax.grid(True, linestyle="--", alpha=0.5)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.markdown("### 📋 Load–Penetration Table")
        st.dataframe(df, use_container_width=True)
        st.image(img_buf)

        c1, c2, c3 = st.columns(3)
        c1.metric("CBR at 2.5 mm", f"{cbr_25:.2f}%")
        c2.metric("CBR at 5.0 mm", f"{cbr_50:.2f}%")
        c3.metric("Final CBR", f"{final_cbr:.2f}%")
        st.info(f"🏷️ **{conclusion}**")

        st.session_state.cbr_res = {
            "df": df, "cbr_25": cbr_25, "cbr_50": cbr_50,
            "final_cbr": final_cbr, "conclusion": conclusion,
            "img_buf": img_buf, "ring_const": ring_const,
        }

        report_buf = _generate_report(
            ring_const, df, cbr_25, cbr_50, final_cbr,
            conclusion, img_buf, PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="CBR_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":        PROCEDURE,
            "formulas":         FORMULAS,
            "data":             df,
            "graph":            img_buf,
            "CBR at 2.5mm (%)": round(cbr_25, 2),
            "CBR at 5.0mm (%)": round(cbr_50, 2),
            "Final CBR (%)":    round(final_cbr, 2),
            "Subgrade Class":   conclusion.split(".")[0],
        }

    return None