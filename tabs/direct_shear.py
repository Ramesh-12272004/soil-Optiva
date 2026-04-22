import streamlit as st
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the shear strength parameters – cohesion (c) and angle of
  internal friction (ϕ) – of soil using the direct shear test as per
  IS 2720 (Part 13) – 1986.

Apparatus Required:
  - Direct shear apparatus (motor-driven with gearbox)
  - Shear box (60 mm × 60 mm or 100 mm × 100 mm square section)
  - Normal load hanger and dead weights
  - Proving ring with dial gauge (to measure shear force)
  - Horizontal deformation dial gauge
  - Vertical deformation (settlement) dial gauge
  - Porous stones and filter paper
  - Weighing balance, spatula, tamping rod

Theory:
  In a direct shear test, the soil specimen is split horizontally. A normal
  load σ_n is applied on the top half; the lower half is driven horizontally
  while the upper half is restrained. The shear force required to cause failure
  is measured. The Mohr-Coulomb failure envelope is obtained by plotting
  σ_n vs τ_max for at least 3 trials.

  τ = c + σ_n × tan(ϕ)

Step-by-Step Procedure:
  1. Measure the shear box dimensions; calculate the cross-sectional area A.
  2. Place soil specimen (undisturbed or remoulded) in the shear box;
     place porous stones above and below.
  3. Apply the normal load σ_n (kg/cm²) and allow consolidation (at least
     24 h for cohesive soils; 5–10 min for sands).
  4. Separate the two halves by 0.5–1 mm; remove the locking screws.
  5. Start the motor at a slow strain rate (0.2 mm/min for cohesive;
     1–2 mm/min for granular soils).
  6. Record proving ring (shear force) and horizontal deformation readings
     at regular intervals until peak shear is passed or 20% strain reached.
  7. Repeat for at least 3 different normal stresses.
  8. Plot τ vs deformation curves and the Mohr-Coulomb envelope.

Precautions:
  - Remove locking pins before starting the motor.
  - Maintain a constant strain rate throughout.
  - Record proving ring to the nearest division at each reading.
  - Ensure the normal load hanger is vertical with no friction.
"""

FORMULAS = """
Cross-sectional Area:
  A = (side_length / 10)²    [cm²]  (side_length in mm → cm)

Shear Force:
  F_s = Proving Ring Reading (div) × Proving Ring Constant (kg/div)

Shear Stress:
  τ = F_s / A    [kg/cm²]

Horizontal Deformation:
  δ = Reading (div) × Dial Gauge Least Count (mm/div)    [mm]

Mohr-Coulomb Failure Criterion:
  τ = c + σ_n × tan(ϕ)

  Best-fit linear regression of τ_max vs σ_n gives:
    slope → tan(ϕ)  ⟹  ϕ = arctan(slope)   [degrees]
    intercept → c   [kg/cm²]

Angle of Friction:
  ϕ = arctan(Δτ / Δσ_n)    [degrees]
"""


def _generate_report(box_dim, prc, dial_lc, area, all_dfs, normal_stresses,
                     shear_stresses, cohesion, phi, img_mohr, procedure, formulas):
    doc = Document()
    doc.add_heading("Direct Shear Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 13) – 1986")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine cohesion (c) and angle of internal friction (ϕ) "
        "using the direct shear test."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Shear Box Side = {box_dim:.1f} mm")
    doc.add_paragraph(f"Proving Ring Constant = {prc:.4f} kg/div")
    doc.add_paragraph(f"Dial Gauge LC = {dial_lc:.4f} mm/div")
    doc.add_paragraph(f"Shear Area A = {area:.4f} cm²")
    doc.add_paragraph(f"Cohesion c = {cohesion:.4f} kg/cm²")
    doc.add_paragraph(f"Angle of Internal Friction ϕ = {phi:.2f}°")

    for t_idx, df_trial in enumerate(all_dfs):
        doc.add_heading(f"Trial {t_idx + 1} – σₙ = {normal_stresses[t_idx]:.3f} kg/cm²", 2)
        tbl = doc.add_table(rows=1, cols=len(df_trial.columns))
        tbl.style = "Table Grid"
        for ci, col in enumerate(df_trial.columns):
            tbl.rows[0].cells[ci].text = col
        for _, row in df_trial.iterrows():
            cells = tbl.add_row().cells
            for ci, val in enumerate(row):
                cells[ci].text = f"{val:.4f}" if isinstance(val, float) else str(val)

    doc.add_heading("Mohr-Coulomb Failure Envelope", 1)
    img_mohr.seek(0)
    doc.add_picture(img_mohr, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"Cohesion c = {cohesion:.4f} kg/cm² and "
        f"Angle of Internal Friction ϕ = {phi:.2f}°. "
        f"{'Dense / overconsolidated soil.' if phi > 30 else 'Loose / normally consolidated soil.'}"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("✂️ Direct Shear Test (IS 2720 Part 13 : 1986)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "ds_res" not in st.session_state:
        st.session_state.ds_res = None

    # ── Apparatus constants ──
    st.markdown("### 🔧 Apparatus Constants")
    c1, c2, c3 = st.columns(3)
    box_dim = c1.number_input("Box Side Length (mm)", value=60.0, min_value=1.0, format="%.1f", key="ds_box")
    prc     = c2.number_input("Proving Ring Constant (kg/div)", value=0.0, min_value=0.0, format="%.4f", key="ds_prc")
    dial_lc = c3.number_input("Dial Gauge LC (mm/div)",        value=0.01, min_value=0.0001, format="%.4f", key="ds_dlc")
    area = (box_dim / 10) ** 2
    st.info(f"📐 Shear Area A = **{area:.3f} cm²**")

    # ── Trials ──
    num_trials = st.number_input("Number of Normal Stress Trials (min 2)", min_value=2, max_value=5, value=3, step=1, key="ds_nt")
    n_readings = st.number_input("Readings per Trial", min_value=3, max_value=20, value=6, step=1, key="ds_nr")

    # Common horizontal deformation column
    st.markdown("### 📏 Horizontal Deformation Readings (dial divisions, same for all trials)")
    if "ds_hdef" not in st.session_state or len(st.session_state.ds_hdef) != n_readings:
        st.session_state.ds_hdef = [0.0] * n_readings
    cols_def = st.columns(min(n_readings, 4))
    for i in range(n_readings):
        st.session_state.ds_hdef[i] = cols_def[i % 4].number_input(
            f"H-Def {i + 1}", value=st.session_state.ds_hdef[i],
            min_value=0.0, format="%.2f", key=f"ds_hd_{i}"
        )

    # Per-trial inputs
    if "ds_trials" not in st.session_state or len(st.session_state.ds_trials) != num_trials:
        st.session_state.ds_trials = [
            {"sigma_n": 0.0, "pr": [0.0] * n_readings}
            for _ in range(num_trials)
        ]
    for t in range(num_trials):
        if len(st.session_state.ds_trials[t]["pr"]) != n_readings:
            st.session_state.ds_trials[t]["pr"] = [0.0] * n_readings

    st.markdown("### 🧪 Trial Data")
    for t in range(num_trials):
        st.markdown(f"#### Trial {t + 1}")
        st.session_state.ds_trials[t]["sigma_n"] = st.number_input(
            f"Normal Stress σₙ (kg/cm²)", value=st.session_state.ds_trials[t]["sigma_n"],
            min_value=0.0, format="%.3f", key=f"ds_sn_{t}"
        )
        pr_cols = st.columns(min(n_readings, 4))
        for i in range(n_readings):
            st.session_state.ds_trials[t]["pr"][i] = pr_cols[i % 4].number_input(
                f"PR {i + 1}", value=st.session_state.ds_trials[t]["pr"][i],
                min_value=0.0, format="%.2f", key=f"ds_pr_{t}_{i}"
            )

    if st.button("🔄 Reset"):
        st.session_state.ds_trials = [{"sigma_n": 0.0, "pr": [0.0] * n_readings} for _ in range(num_trials)]
        st.session_state.ds_hdef   = [0.0] * n_readings
        st.session_state.ds_res    = None
        st.rerun()

    if st.button("📊 Calculate"):
        normal_stresses = []
        shear_stresses  = []
        all_dfs         = []
        figs_trial      = []

        for t, trial in enumerate(st.session_state.ds_trials[:num_trials]):
            sigma_n = trial["sigma_n"]
            pr      = trial["pr"]

            h_defl  = [st.session_state.ds_hdef[i] * dial_lc for i in range(n_readings)]
            F_s     = [pr[i] * prc for i in range(n_readings)]
            tau     = [f / area if area > 0 else 0 for f in F_s]

            df = pd.DataFrame({
                "H-Defl Reading (div)":   st.session_state.ds_hdef[:n_readings],
                "Deformation δ (mm)":     [round(x, 4) for x in h_defl],
                "PR Reading (div)":       pr[:n_readings],
                "Shear Force F_s (kg)":   [round(x, 4) for x in F_s],
                "Shear Stress τ (kg/cm²)":[round(x, 5) for x in tau],
            })

            tau_max = max(tau)
            normal_stresses.append(sigma_n)
            shear_stresses.append(tau_max)
            all_dfs.append(df)

            st.markdown(f"#### Trial {t + 1} – σₙ = {sigma_n:.3f} kg/cm²")
            st.dataframe(df, use_container_width=True)

            fig_t, ax_t = plt.subplots(figsize=(6, 3.5))
            ax_t.plot(df["Deformation δ (mm)"], df["Shear Stress τ (kg/cm²)"],
                      marker="o", color="#0a68cc", linewidth=2)
            ax_t.axhline(tau_max, color="red", linestyle="--", linewidth=1.2,
                         label=f"τ_max = {tau_max:.4f}")
            ax_t.set_xlabel("Deformation δ (mm)")
            ax_t.set_ylabel("Shear Stress τ (kg/cm²)")
            ax_t.set_title(f"Trial {t + 1}: τ vs δ  (σₙ = {sigma_n:.3f} kg/cm²)")
            ax_t.legend()
            ax_t.grid(True, linestyle="--", alpha=0.5)
            st.pyplot(fig_t)
            buf_t = BytesIO()
            fig_t.savefig(buf_t, format="png", dpi=120, bbox_inches="tight")
            buf_t.seek(0)
            figs_trial.append(buf_t)
            plt.close(fig_t)

        if len(normal_stresses) >= 2:
            coeffs  = np.polyfit(normal_stresses, shear_stresses, 1)
            phi_deg = math.degrees(math.atan(coeffs[0]))
            cohesion= coeffs[1]

            st.success(f"**Cohesion c = {cohesion:.4f} kg/cm²**")
            st.success(f"**Angle of Internal Friction ϕ = {phi_deg:.2f}°**")

            x_line = np.linspace(0, max(normal_stresses) * 1.2, 100)
            y_line = coeffs[0] * x_line + coeffs[1]

            fig_m, ax_m = plt.subplots(figsize=(7, 4))
            ax_m.scatter(normal_stresses, shear_stresses,
                         color="#0a68cc", s=80, zorder=5, label="τ_max points")
            ax_m.plot(x_line, y_line, color="#e05c00", linewidth=2,
                      label=f"τ = {cohesion:.3f} + σₙ·tan({phi_deg:.1f}°)")
            ax_m.set_xlabel("Normal Stress σₙ (kg/cm²)")
            ax_m.set_ylabel("Max Shear Stress τ_max (kg/cm²)")
            ax_m.set_title("Mohr-Coulomb Failure Envelope")
            ax_m.legend()
            ax_m.grid(True, linestyle="--", alpha=0.5)
            st.pyplot(fig_m)

            img_mohr = BytesIO()
            fig_m.savefig(img_mohr, format="png", dpi=150, bbox_inches="tight")
            img_mohr.seek(0)
            plt.close(fig_m)

            st.session_state.ds_res = {
                "box_dim": box_dim, "prc": prc, "dial_lc": dial_lc, "area": area,
                "all_dfs": all_dfs, "normal_stresses": normal_stresses,
                "shear_stresses": shear_stresses, "cohesion": cohesion,
                "phi": phi_deg, "img_mohr": img_mohr,
            }

    if st.session_state.ds_res is not None:
        res = st.session_state.ds_res
        report_buf = _generate_report(
            res["box_dim"], res["prc"], res["dial_lc"], res["area"],
            res["all_dfs"], res["normal_stresses"], res["shear_stresses"],
            res["cohesion"], res["phi"], res["img_mohr"],
            PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Direct_Shear_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":                    PROCEDURE,
            "formulas":                     FORMULAS,
            "graph":                        res["img_mohr"],
            "Cohesion c (kg/cm2)":          round(res["cohesion"], 4),
            "Friction Angle phi (degrees)": round(res["phi"], 2),
        }

    return None