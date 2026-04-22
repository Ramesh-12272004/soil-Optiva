import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the Plastic Limit (PL) of fine-grained soil as per
  IS 2720 (Part 5) – 1985.

Apparatus Required:
  - Glass plate (frosted / non-absorbent)
  - Weighing balance (accuracy 0.01 g)
  - Moisture content cans with lids
  - Oven (105 °C – 110 °C)
  - IS sieve 425 µm
  - Wash bottle with distilled water
  - Rod of 3 mm diameter (for reference)

Theory:
  The Plastic Limit is the lowest water content at which soil can be
  rolled into a 3 mm diameter thread without crumbling. It marks the
  boundary between the plastic and semi-solid states.

  Plasticity Index (PI) = LL – PL
  PI indicates the range of water content over which soil behaves plastically.

Step-by-Step Procedure:
  1. Pass air-dried soil through a 425 µm sieve; take about 20 g.
  2. Mix thoroughly with distilled water until it becomes plastic.
  3. Roll a small ball (~6 g) between the palms to a uniform ball.
  4. Place the ball on the glass plate and roll with fingers to a
     uniform thread of 3 mm diameter.
  5. If the thread does not crumble, fold it, re-roll, and repeat
     until it just crumbles at exactly 3 mm diameter.
  6. Collect the crumbled pieces immediately into a moisture can.
  7. Determine the water content of the crumbled thread.
  8. Repeat with fresh soil (minimum 3 trials).
  9. The average water content at crumbling = Plastic Limit.

Precautions:
  - Rolling must be uniform, using the full length of the fingers.
  - Crumbling must occur AT 3 mm, not before or after.
  - Avoid excessive drying during rolling.
  - At least 3 consistent trials are required.
"""

FORMULAS = """
Water Content per trial (w%):
  w (%) = [(W2 - W3) / (W3 - W1)] × 100

  Where:
    W1 = Mass of empty moisture can (g)
    W2 = Mass of can + wet soil     (g)
    W3 = Mass of can + dry soil     (g)

Plastic Limit (PL):
  PL = Average of water contents from all valid trials (%)

Plasticity Index (PI):
  PI = Liquid Limit (LL) – Plastic Limit (PL)

Activity (A):
  A = PI / (% clay fraction)  [if clay % is known]

IS Classification based on PI:
  PI < 7   → Non-plastic to slightly plastic
  7 ≤ PI < 17 → Moderately plastic
  PI ≥ 17  → Highly plastic
"""


def _calc_wc(w1, w2, w3):
    if w2 <= w3 or w3 <= w1:
        return np.nan
    return (w2 - w3) / (w3 - w1) * 100


def _generate_report(df, pl, ll_input, pi, procedure, formulas):
    doc = Document()
    doc.add_heading("Plastic Limit Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 5) – 1985")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the Plastic Limit (PL) of fine-grained soil and "
        "calculate the Plasticity Index (PI)."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Observation Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(round(val, 3) if isinstance(val, float) else val)

    doc.add_heading("Results", 1)
    doc.add_paragraph(f"Plastic Limit (PL)       = {pl:.2f} %")
    if ll_input > 0:
        doc.add_paragraph(f"Liquid Limit (LL)        = {ll_input:.2f} %")
        doc.add_paragraph(f"Plasticity Index (PI)    = {pi:.2f} %")

    doc.add_heading("Conclusion", 1)
    if ll_input > 0:
        if pi < 7:
            plas = "non-plastic to slightly plastic"
        elif pi < 17:
            plas = "moderately plastic"
        else:
            plas = "highly plastic"
        doc.add_paragraph(
            f"The Plastic Limit is {pl:.2f}% and the Plasticity Index is {pi:.2f}%, "
            f"indicating a {plas} soil."
        )
    else:
        doc.add_paragraph(
            f"The Plastic Limit of the soil sample is {pl:.2f}%."
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🟤 Plastic Limit Test (IS 2720 Part 5 : 1985)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    # ── Session state ──
    if "pl_num_trials" not in st.session_state:
        st.session_state.pl_num_trials = 3
    if "pl_inputs" not in st.session_state:
        st.session_state.pl_inputs = [
            {"w1": 0.0, "w2": 0.0, "w3": 0.0}
            for _ in range(3)
        ]
    if "pl_ll" not in st.session_state:
        st.session_state.pl_ll = 0.0

    num_trials = st.number_input(
        "Number of Trials (min 3)", min_value=3, max_value=10,
        value=st.session_state.pl_num_trials, step=1
    )
    while len(st.session_state.pl_inputs) < num_trials:
        st.session_state.pl_inputs.append({"w1": 0.0, "w2": 0.0, "w3": 0.0})
    st.session_state.pl_inputs = st.session_state.pl_inputs[:num_trials]
    st.session_state.pl_num_trials = num_trials

    st.markdown("### ⚖️ Moisture Content Readings")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3 = st.columns(3)
        inp = st.session_state.pl_inputs[i]
        inp["w1"] = c1.number_input(f"W1 – Empty Can (g)", value=inp["w1"],
                                    min_value=0.0, format="%.3f", key=f"pl_w1_{i}")
        inp["w2"] = c2.number_input(f"W2 – Wet Soil+Can (g)", value=inp["w2"],
                                    min_value=0.0, format="%.3f", key=f"pl_w2_{i}")
        inp["w3"] = c3.number_input(f"W3 – Dry Soil+Can (g)", value=inp["w3"],
                                    min_value=0.0, format="%.3f", key=f"pl_w3_{i}")

        wc = _calc_wc(inp["w1"], inp["w2"], inp["w3"])
        if np.isnan(wc):
            if not (inp["w1"] == inp["w2"] == inp["w3"] == 0.0):
                st.warning(f"Trial {i + 1}: Invalid weights.")
        else:
            st.info(f"Trial {i + 1} Water Content = **{wc:.2f}%**")

    st.markdown("---")
    st.markdown("### 📌 Optional – Enter Liquid Limit for PI Calculation")
    st.session_state.pl_ll = st.number_input(
        "Liquid Limit LL (%)", value=st.session_state.pl_ll,
        min_value=0.0, format="%.2f", key="pl_ll_input"
    )

    if st.button("🔄 Reset All Inputs"):
        st.session_state.pl_inputs = [{"w1": 0.0, "w2": 0.0, "w3": 0.0} for _ in range(num_trials)]
        st.session_state.pl_ll = 0.0
        st.rerun()

    if st.button("📊 Calculate Plastic Limit"):
        rows = []
        wc_list = []
        for i, inp in enumerate(st.session_state.pl_inputs[:num_trials]):
            wc = _calc_wc(inp["w1"], inp["w2"], inp["w3"])
            rows.append({
                "Trial":             i + 1,
                "W1 – Empty Can (g)":     inp["w1"],
                "W2 – Wet Soil+Can (g)":  inp["w2"],
                "W3 – Dry Soil+Can (g)":  inp["w3"],
                "Water Content (%)":  round(wc, 2) if not np.isnan(wc) else np.nan,
            })
            if not np.isnan(wc):
                wc_list.append(wc)

        df = pd.DataFrame(rows)

        if len(wc_list) < 2:
            st.error("At least 2 valid trials required.")
            return None

        pl   = float(np.mean(wc_list))
        ll   = float(st.session_state.pl_ll)
        pi   = ll - pl if ll > 0 else None

        st.markdown("### 📊 Observation Table")
        st.dataframe(df.round(3), use_container_width=True)

        c1, c2, c3 = st.columns(3)
        c1.metric("Plastic Limit (PL)", f"{pl:.2f}%")
        if ll > 0 and pi is not None:
            c2.metric("Liquid Limit (LL)", f"{ll:.2f}%")
            c3.metric("Plasticity Index (PI)", f"{pi:.2f}%")

            if pi < 7:
                plas_desc = "Non-plastic to Slightly Plastic"
                st.info(f"🏷️ Plasticity: **{plas_desc}** (PI = {pi:.2f}%)")
            elif pi < 17:
                plas_desc = "Moderately Plastic"
                st.warning(f"🏷️ Plasticity: **{plas_desc}** (PI = {pi:.2f}%)")
            else:
                plas_desc = "Highly Plastic"
                st.error(f"🏷️ Plasticity: **{plas_desc}** (PI = {pi:.2f}%)")
        else:
            pi = 0.0
            plas_desc = "N/A (LL not provided)"
            st.info(f"Plastic Limit = {pl:.2f}%. Enter LL above to get PI.")

        # Word report
        report_buf = _generate_report(df, pl, ll, pi if pi else 0.0, PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Plastic_Limit_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":            PROCEDURE,
            "formulas":             FORMULAS,
            "data":                 df,
            "Plastic Limit PL (%)": round(pl, 2),
            "Liquid Limit LL (%)":  round(ll, 2) if ll > 0 else None,
            "Plasticity Index PI (%)": round(pi, 2) if pi else None,
            "Plasticity Class":     plas_desc,
        }

    return None