import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from docx import Document
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the specific gravity (G) of soil solids using the density
  bottle (pycnometer) method as per IS 2720 (Part 3, Sec 1) – 1980.

Apparatus Required:
  - Density bottle (50 mL capacity) with ground-glass stopper
  - Vacuum pump / de-airing facility
  - Weighing balance (accuracy 0.001 g)
  - Oven (105 °C – 110 °C)
  - Distilled water or kerosene / CCl₄ (for soils with fine organic matter)
  - Thermometer
  - Wash bottle

Theory:
  Specific Gravity is the ratio of the mass of a given volume of soil
  solids to the mass of an equal volume of distilled water at 4 °C.
  Typical values:
    Sand / inorganic soil  : 2.60 – 2.68
    Silty clay             : 2.67 – 2.75
    Dense clay / minerals  : 2.75 – 2.85
    Organic soil           : < 2.60

Step-by-Step Procedure:
  1. Clean and oven-dry the density bottle; cool in desiccator. Record W1.
  2. Add 10–15 g of oven-dried soil; stopper and weigh. Record W2.
  3. Add the pore fluid (distilled water or CCl₄) to half-fill the bottle;
     remove stopper and stir. Place on vacuum pump for ≥ 1 hour to de-air.
  4. Fill completely to the brim with the same pore fluid (no air bubbles).
     Replace stopper; wipe dry; weigh. Record W3.
  5. Empty, clean, and fill bottle completely with pore fluid only; stopper,
     wipe, and weigh. Record W4.
  6. Repeat steps 1–5 for at least 3 trials.

Precautions:
  - All weighings must be done at the same temperature.
  - Ensure no air bubbles remain during de-airing.
  - Use the same pore fluid throughout a single trial.
  - Wipe the outside of the bottle thoroughly before each weighing.
"""

FORMULAS = """
Density of pore fluid (Gc):
  Gc = (W4 – W1) / V

Specific Gravity (G):
  G = (W2 – W1) × Gc / [(W4 – W1) – (W3 – W2)]

Simplified (when Gc ≈ 1, i.e. distilled water at ≈ 27 °C):
  G ≈ (W2 – W1) / [(W2 – W1) – (W3 – W4)]

Where:
  W1 = Mass of empty density bottle (g)
  W2 = Mass of bottle + dry soil (g)
  W3 = Mass of bottle + soil + pore fluid (g)
  W4 = Mass of bottle + pore fluid only (g)
  V  = Internal volume of density bottle (cm³)
"""


def _generate_report(df, avg_g, soil_type, volume, procedure, formulas):
    doc = Document()
    doc.add_heading("Specific Gravity Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 3, Sec 1) – 1980")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph("To determine the specific gravity of soil solids using the density bottle method.")

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Density Bottle Volume (V): {volume:.2f} cm³")

    doc.add_heading("Observation & Results Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(round(val, 4) if isinstance(val, float) else val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Average Specific Gravity G = {avg_g:.3f}")
    doc.add_paragraph(f"Soil Type (interpretation): {soil_type}")

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The specific gravity of the given soil is {avg_g:.3f}, "
        f"which corresponds to {soil_type}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("⚗️ Specific Gravity – Density Bottle Method (IS 2720 Part 3 Sec 1 : 1980)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "sg_res" not in st.session_state:
        st.session_state.sg_res = None

    # ── Inputs ──
    st.markdown("### 🔢 Density Bottle Volume")
    volume = st.number_input("Volume V (cm³)", value=50.0, min_value=1.0, format="%.2f", key="sg_vol")

    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1, key="sg_nt")

    if "sg_inputs" not in st.session_state or len(st.session_state.sg_inputs) != num_trials:
        st.session_state.sg_inputs = [{"w1": 0.0, "w2": 0.0, "w3": 0.0, "w4": 0.0}
                                      for _ in range(num_trials)]

    st.markdown("### ⚖️ Enter Weights for Each Trial")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3, c4 = st.columns(4)
        inp = st.session_state.sg_inputs[i]
        inp["w1"] = c1.number_input(f"W1 – Empty Bottle (g)",        value=inp["w1"], min_value=0.0, format="%.3f", key=f"sg_w1_{i}")
        inp["w2"] = c2.number_input(f"W2 – Bottle + Dry Soil (g)",   value=inp["w2"], min_value=0.0, format="%.3f", key=f"sg_w2_{i}")
        inp["w3"] = c3.number_input(f"W3 – Bottle + Soil + Fluid (g)", value=inp["w3"], min_value=0.0, format="%.3f", key=f"sg_w3_{i}")
        inp["w4"] = c4.number_input(f"W4 – Bottle + Fluid Only (g)", value=inp["w4"], min_value=0.0, format="%.3f", key=f"sg_w4_{i}")

    if st.button("🔄 Reset"):
        st.session_state.sg_inputs = [{"w1": 0.0, "w2": 0.0, "w3": 0.0, "w4": 0.0} for _ in range(num_trials)]
        st.session_state.sg_res = None
        st.rerun()

    if st.button("📊 Calculate Specific Gravity"):
        rows = []
        g_list = []
        for i, inp in enumerate(st.session_state.sg_inputs[:num_trials]):
            w1, w2, w3, w4 = inp["w1"], inp["w2"], inp["w3"], inp["w4"]
            if not (w2 > w1 and w4 > w1):
                rows.append({"Trial": i + 1, "W1 (g)": w1, "W2 (g)": w2,
                             "W3 (g)": w3, "W4 (g)": w4, "Gc": None, "G": None})
                continue
            Gc   = (w4 - w1) / volume
            denom = (w4 - w1) - (w3 - w2)
            if denom <= 0:
                rows.append({"Trial": i + 1, "W1 (g)": w1, "W2 (g)": w2,
                             "W3 (g)": w3, "W4 (g)": w4, "Gc": round(Gc, 4), "G": None})
                continue
            G = (w2 - w1) * Gc / denom
            g_list.append(G)
            rows.append({"Trial": i + 1, "W1 (g)": w1, "W2 (g)": w2,
                         "W3 (g)": w3, "W4 (g)": w4, "Gc": round(Gc, 4), "G": round(G, 4)})

        df = pd.DataFrame(rows)

        if not g_list:
            st.error("No valid trials. Check that W2 > W1, W4 > W1, and denom > 0.")
            return None

        avg_g = float(np.mean(g_list))

        if avg_g < 2.55:
            soil_type = "Organic / peat soil"
        elif avg_g < 2.67:
            soil_type = "Sand / inorganic soil"
        elif avg_g <= 2.75:
            soil_type = "Silty clay"
        else:
            soil_type = "Dense clay / heavy minerals"

        st.session_state.sg_res = {
            "df": df, "avg_g": avg_g, "soil_type": soil_type, "volume": volume
        }

    if st.session_state.sg_res is not None:
        res = st.session_state.sg_res
        st.markdown("### 📋 Results Table")
        st.dataframe(res["df"].round(4), use_container_width=True)

        c1, c2 = st.columns(2)
        c1.metric("Average Specific Gravity G", f"{res['avg_g']:.3f}")
        c2.metric("Soil Type", res["soil_type"])

        report_buf = _generate_report(
            res["df"], res["avg_g"], res["soil_type"],
            res["volume"], PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Specific_Gravity_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":                 PROCEDURE,
            "formulas":                  FORMULAS,
            "data":                      res["df"],
            "Average Specific Gravity G": round(res["avg_g"], 3),
            "Soil Type":                 res["soil_type"],
        }

    return None