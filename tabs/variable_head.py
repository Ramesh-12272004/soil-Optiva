import streamlit as st
import math
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the coefficient of permeability (k) of fine-grained
  soils using the falling (variable) head apparatus as per IS 2720
  (Part 17) – 1986.

Apparatus Required:
  - Variable head (falling head) permeameter
  - Graduated standpipe (burette) – cross-section area 'a'
  - Stopwatch
  - Steel scale / vernier calliper
  - Weighing balance
  - Filter paper and porous stones

Theory:
  In a falling head test the head driving flow decreases over time.
  The standpipe head drops from h₁ to h₂ in time t. By recording the
  head at the start (h₁) and end (h₂) and integrating Darcy's law
  over the time interval, k is derived.

  Typical k ranges for fine-grained soils:
    10⁻⁵ – 10⁻³ cm/s  → Silt / fine sand
    10⁻⁷ – 10⁻⁵ cm/s  → Sandy clay / silty clay
    < 10⁻⁷ cm/s       → Clay

Step-by-Step Procedure:
  1. Measure the cross-section of the standpipe (a) and the specimen
     dimensions (A = cross-section area, L = length).
  2. Saturate the specimen by allowing upward flow for ≥ 30 min.
  3. Fill the standpipe to a known initial head h₁ (cm) above the
     datum (overflow level).
  4. Start the stopwatch when the water level in the standpipe
     passes a reference mark.
  5. Stop the watch when the level drops to a second reference mark h₂.
  6. Record h₁, h₂, and elapsed time t.
  7. Repeat at least 3 times; h₁ can be the same or different.

Precautions:
  - Ensure full saturation before measurements.
  - Measure h₁ and h₂ from the same datum (typically the outlet level).
  - Avoid temperature fluctuations during the test.
  - De-air the inlet and standpipe water.
"""

FORMULAS = """
Coefficient of Permeability (k):
  k = (2.303 × a × L) / (A × t) × log₁₀(h₁ / h₂)    [cm/s]

Where:
  a  = Cross-sectional area of standpipe (cm²)
  L  = Length of soil specimen (cm)
  A  = Cross-sectional area of soil specimen (cm²)
  t  = Elapsed time (s)
  h₁ = Initial head (cm)
  h₂ = Final head (cm)

Conversion:
  k (m/s) = k (cm/s) × 0.01
"""


def _classify(k_cms):
    if k_cms > 1e-3:
        return "Silty Sand / Fine Sand"
    elif k_cms > 1e-5:
        return "Silt / Sandy Clay"
    elif k_cms > 1e-7:
        return "Silty Clay"
    else:
        return "Clay"


def _generate_report(a, A, L, df, avg_k, soil_type, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("Variable Head (Falling Head) Permeability Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 17) – 1986")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the coefficient of permeability k of fine-grained soil "
        "using the falling head permeameter."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Standpipe area a = {a:.4f} cm²")
    doc.add_paragraph(f"Specimen area  A = {A:.4f} cm²")
    doc.add_paragraph(f"Specimen length L = {L:.3f} cm")

    doc.add_heading("Results Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.6f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Average k = {avg_k:.6f} cm/s  ({avg_k * 0.01:.6e} m/s)")
    doc.add_paragraph(f"Soil Classification: {soil_type}")

    doc.add_heading("Graph", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The coefficient of permeability is {avg_k:.6f} cm/s, "
        f"indicating the soil is classified as {soil_type}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🌊 Variable Head (Falling Head) Permeability Test (IS 2720 Part 17 : 1986)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "vh_res" not in st.session_state:
        st.session_state.vh_res = None

    # ── Constants ──
    st.markdown("### 📏 Specimen & Standpipe Dimensions")
    c1, c2, c3 = st.columns(3)
    a = c1.number_input("Standpipe area a (cm²)", value=1.00, min_value=0.001, format="%.4f", key="vh_a")
    A = c2.number_input("Specimen area  A (cm²)", value=50.00, min_value=0.001, format="%.4f", key="vh_A")
    L = c3.number_input("Specimen length L (cm)", value=10.00, min_value=0.001, format="%.3f", key="vh_L")

    # ── Trial inputs ──
    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1, key="vh_nt")

    if "vh_inputs" not in st.session_state or len(st.session_state.vh_inputs) != num_trials:
        st.session_state.vh_inputs = [{"h1": 0.0, "h2": 0.0, "t": 0.0} for _ in range(num_trials)]

    st.markdown("### 📋 Enter Trial Data")
    for i in range(num_trials):
        col1, col2, col3 = st.columns(3)
        inp = st.session_state.vh_inputs[i]
        inp["h1"] = col1.number_input(f"h₁ – Initial Head (cm) [T{i+1}]", value=inp["h1"], min_value=0.0, format="%.3f", key=f"vh_h1_{i}")
        inp["h2"] = col2.number_input(f"h₂ – Final Head (cm)   [T{i+1}]", value=inp["h2"], min_value=0.0, format="%.3f", key=f"vh_h2_{i}")
        inp["t"]  = col3.number_input(f"t – Elapsed Time (s)   [T{i+1}]", value=inp["t"],  min_value=0.0, format="%.3f", key=f"vh_t_{i}")

    if st.button("🔄 Reset"):
        st.session_state.vh_inputs = [{"h1": 0.0, "h2": 0.0, "t": 0.0} for _ in range(num_trials)]
        st.session_state.vh_res = None
        st.rerun()

    if st.button("📊 Calculate"):
        rows = []
        for i, inp in enumerate(st.session_state.vh_inputs[:num_trials]):
            h1, h2, t = inp["h1"], inp["h2"], inp["t"]
            if h1 > h2 > 0 and t > 0 and a > 0 and A > 0 and L > 0:
                k = (2.303 * a * L) / (A * t) * math.log10(h1 / h2)
            else:
                k = 0.0
            rows.append({
                "Trial":    i + 1,
                "h₁ (cm)":  h1,
                "h₂ (cm)":  h2,
                "t (s)":    t,
                "k (cm/s)": round(k, 7),
                "k (m/s)":  round(k * 0.01, 9),
            })

        df = pd.DataFrame(rows)
        valid_k = df[df["k (cm/s)"] > 0]["k (cm/s)"]

        if valid_k.empty:
            st.error("No valid trials. Ensure h₁ > h₂ > 0 and t > 0.")
            return None

        avg_k     = float(valid_k.mean())
        soil_type = _classify(avg_k)

        # Plot
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.plot(df["Trial"], df["k (cm/s)"], marker="o", color="#0a68cc", linewidth=2)
        ax.axhline(avg_k, color="#e05c00", linestyle="--", linewidth=1.5,
                   label=f"Avg k = {avg_k:.6f} cm/s")
        ax.set_xlabel("Trial Number")
        ax.set_ylabel("k (cm/s)")
        ax.set_title("Falling Head – Permeability per Trial")
        ax.legend()
        ax.grid(True, linestyle="--", alpha=0.5)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.session_state.vh_res = {
            "a": a, "A": A, "L": L,
            "df": df, "avg_k": avg_k,
            "soil_type": soil_type, "img_buf": img_buf
        }

    if st.session_state.vh_res is not None:
        res = st.session_state.vh_res
        st.markdown("### 📋 Results Table")
        st.dataframe(res["df"].style.format(precision=7), use_container_width=True)

        c1, c2 = st.columns(2)
        c1.metric("Average k (cm/s)", f"{res['avg_k']:.6f}")
        c2.metric("Soil Classification", res["soil_type"])
        st.image(res["img_buf"])

        report_buf = _generate_report(
            res["a"], res["A"], res["L"],
            res["df"], res["avg_k"], res["soil_type"],
            res["img_buf"], PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Variable_Head_Permeability_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":           PROCEDURE,
            "formulas":            FORMULAS,
            "data":                res["df"],
            "graph":               res["img_buf"],
            "Average k (cm/s)":    round(res["avg_k"], 7),
            "Soil Classification": res["soil_type"],
        }

    return None