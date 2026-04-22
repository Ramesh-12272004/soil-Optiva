import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def run():

    st.subheader("Sieve Analysis (IS 2720 Part 4)")

    # -----------------------------
    # PROCEDURE & FORMULAS
    # -----------------------------
    procedure_text = """
🎯 Objective:
To determine the particle size distribution of coarse-grained soil using a standard set of sieves
as per IS 2720 (Part 4) – 1985.

🧪 Apparatus Required:
- Standard IS sieves: 4.75 mm, 2.36 mm, 1.18 mm, 0.6 mm, 0.425 mm, 0.3 mm, 0.15 mm, 0.075 mm
- Mechanical sieve shaker
- Weighing balance (accuracy 0.1 g)
- Oven (105°C – 110°C)
- Brush & collection tray

📝 Theory:
Sieve analysis, also called gradation testing, is the practice of assessing the particle size
distribution of a granular material. The results are used to classify the soil (IS or USCS system),
compute uniformity coefficient (Cu) and coefficient of curvature (Cc), and evaluate suitability
for various construction uses.

📋 Step-by-Step Procedure:
1. Take a representative oven-dried soil sample (minimum 500 g for coarse soil).
2. Record the total dry weight of the sample.
3. Arrange IS sieves in descending order of aperture size (largest on top).
4. Place the weighed soil into the topmost sieve and fix the cover plate.
5. Mount the sieve stack on the mechanical shaker and operate for 10–15 minutes.
6. Carefully transfer each sieve's retained material to the weighing balance and record the weight.
7. Compute % retained, cumulative % retained, and % finer for each sieve size.
8. Plot the grain size distribution on a semi-logarithmic graph (particle size on log x-axis, % finer on linear y-axis).
9. Read off D10, D30, and D60 from the curve.
10. Compute Cu and Cc and classify the soil.

⚠️ Important Precautions:
- Ensure sieves are clean and undamaged before use.
- Handle sieves gently to avoid particle loss.
- Brush sieves after each test to clear clogged particles.
- Verify that the total weight of retained material + pan weight ≈ original sample weight (loss < 2%).
"""

    formulas_text = """
📐 Formulas Used:

% Retained on each sieve  = (Weight Retained / Total Weight) × 100

Cumulative % Retained      = Running sum of % Retained (from top sieve down)

% Finer (Passing)          = 100 − Cumulative % Retained

Coefficient of Uniformity:
   Cu = D60 / D10

Coefficient of Curvature:
   Cc = (D30²) / (D60 × D10)

Where:
  D10 = Particle size at 10% passing — Effective size
  D30 = Particle size at 30% passing
  D60 = Particle size at 60% passing — Controlling size

Grading Criteria (IS Classification):
  Well-graded gravel (GW): Cu ≥ 4  AND  1 ≤ Cc ≤ 3
  Well-graded sand  (SW): Cu ≥ 6  AND  1 ≤ Cc ≤ 3
  Otherwise: Poorly graded (GP / SP)
"""

    st.markdown("## 📘 Test Procedure")
    with st.expander("📖 Click to read full procedure", expanded=False):
        st.markdown(procedure_text)

    st.markdown("## 📐 Formulas")
    with st.expander("🔢 Click to view formulas", expanded=False):
        st.markdown(formulas_text)

    # -----------------------------
    # INPUT WEIGHTS
    # -----------------------------
    sieve_sizes  = [4.75, 2.36, 1.18, 0.6, 0.425, 0.3, 0.15, 0.075, 0.0]
    sieve_labels = [str(s) if s != 0.0 else "Pan" for s in sieve_sizes]

    if "sieve_weights" not in st.session_state:
        st.session_state.sieve_weights = [0.0] * len(sieve_sizes)

    st.markdown("## 📥 Enter Weight Retained (grams)")
    cols = st.columns(2)
    weights_input = []
    half = len(sieve_labels) // 2
    for i, label in enumerate(sieve_labels):
        with (cols[0] if i < half else cols[1]):
            weight = st.number_input(
                f"Weight Retained on {label} mm",
                min_value=0.0,
                step=0.1,
                value=float(st.session_state.sieve_weights[i]),
                key=f"sieve_input_{i}"
            )
            weights_input.append(weight)

    st.session_state.sieve_weights = weights_input

    # -----------------------------
    # CALCULATION
    # -----------------------------
    if st.button("🔍 Calculate Sieve Analysis"):
        total_weight = sum(weights_input)
        if total_weight == 0:
            st.error("Please enter valid weights.")
            return None

        percent_retained    = [(w / total_weight) * 100 for w in weights_input]
        cumulative_retained = pd.Series(percent_retained).cumsum()
        percent_passing     = 100 - cumulative_retained
        percent_passing.iloc[-1] = 0

        results_df = pd.DataFrame({
            "Sieve Size (mm)":       sieve_labels,
            "Weight Retained (g)":   weights_input,
            "% Retained":            percent_retained,
            "Cumulative % Retained": cumulative_retained,
            "% Passing":             percent_passing,
        })

        results_display = results_df.copy()
        for col in ["Weight Retained (g)", "% Retained", "Cumulative % Retained", "% Passing"]:
            results_display[col] = results_display[col].round(2)

        st.success("Calculation Completed ✅")
        st.dataframe(results_display, use_container_width=True)

        # ---- Grain Size Curve ----
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.semilogx(sieve_sizes[:-1], list(percent_passing)[:-1], marker="o",
                    color="#0a68cc", linewidth=2)
        ax.set_xlabel("Sieve Size (mm)")
        ax.set_ylabel("% Finer")
        ax.set_title("Grain Size Distribution Curve")
        ax.grid(True, which="both", linestyle="--", alpha=0.6)
        ax.invert_xaxis()
        ax.set_ylim(0, 100)
        st.pyplot(fig)

        graph_buffer = BytesIO()
        fig.savefig(graph_buffer, format="PNG", dpi=150, bbox_inches="tight")
        graph_buffer.seek(0)
        plt.close(fig)

        # ---- D10, D30, D60 ----
        def interpolate(p, s, target):
            for j in range(1, len(p)):
                if p[j-1] >= target >= p[j]:
                    x1, y1 = s[j-1], p[j-1]
                    x2, y2 = s[j],   p[j]
                    return x1 + (target - y1) * (x2 - x1) / (y2 - y1)
            return None

        passing_list = list(percent_passing)
        D10 = interpolate(passing_list, sieve_sizes, 10)
        D30 = interpolate(passing_list, sieve_sizes, 30)
        D60 = interpolate(passing_list, sieve_sizes, 60)
        Cu  = (D60 / D10)          if D10 and D60           else None
        Cc  = (D30**2) / (D60*D10) if D10 and D30 and D60   else None

        st.markdown("### 📊 Soil Properties")
        c1, c2, c3, c4, c5 = st.columns(5)
        if D10: c1.metric("D10 (mm)", f"{D10:.3f}")
        if D30: c2.metric("D30 (mm)", f"{D30:.3f}")
        if D60: c3.metric("D60 (mm)", f"{D60:.3f}")
        if Cu:  c4.metric("Cu", f"{Cu:.2f}")
        if Cc:  c5.metric("Cc", f"{Cc:.2f}")

        conclusion    = "Well graded soil." if Cu and Cu >= 5 else "Poorly graded soil."
        suggested_use = (
            "According to IS 2720, this well-graded soil is suitable for road base, sub-base, and general foundation works."
            if Cu and Cu >= 5 else
            "According to IS 2720, this poorly graded soil is more suitable for embankments or fill material but may require stabilization for structural foundations."
        )

        st.markdown("### ✅ Conclusion")
        st.write(conclusion)
        st.markdown("### 💡 Suggested Use")
        st.write(suggested_use)

        # ---- WORD REPORT ----
        doc = Document()
        doc.add_heading("ANITS – Soil Sieve Analysis Report", 0)
        doc.add_heading("Reference Standard: IS 2720 (Part 4) – 1985", level=2)

        doc.add_heading("Test Procedure", level=1)
        for ln in procedure_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())

        doc.add_heading("Formulas", level=1)
        for ln in formulas_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())

        doc.add_page_break()
        doc.add_heading("Results", level=1)
        table = doc.add_table(rows=1, cols=len(results_display.columns))
        table.style = "Table Grid"
        for i, col in enumerate(results_display.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in results_display.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

        doc.add_paragraph("")
        doc.add_heading("Grain Size Distribution Curve", level=1)
        graph_buffer.seek(0)
        doc.add_picture(graph_buffer, width=Inches(5))

        doc.add_heading("Soil Properties", level=1)
        if D10: doc.add_paragraph(f"D10 = {D10:.3f} mm")
        if D30: doc.add_paragraph(f"D30 = {D30:.3f} mm")
        if D60: doc.add_paragraph(f"D60 = {D60:.3f} mm")
        if Cu:  doc.add_paragraph(f"Cu  = {Cu:.2f}")
        if Cc:  doc.add_paragraph(f"Cc  = {Cc:.2f}")

        doc.add_heading("Conclusion",    level=1); doc.add_paragraph(conclusion)
        doc.add_heading("Suggested Use", level=1); doc.add_paragraph(suggested_use)

        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        st.download_button(
            label="⬇️ Download Word Report",
            data=word_buffer,
            file_name="Sieve_Analysis_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # -------------------------------------------------------
        # RETURN RESULT DICT
        # -------------------------------------------------------
        return {
            "procedure":          procedure_text,
            "formulas":           formulas_text,
            "data":               results_df,
            "graph":              graph_buffer,
            "D10 (mm)":           round(D10, 3) if D10 else None,
            "D30 (mm)":           round(D30, 3) if D30 else None,
            "D60 (mm)":           round(D60, 3) if D60 else None,
            "Cu":                 round(Cu,  2)  if Cu  else None,
            "Cc":                 round(Cc,  2)  if Cc  else None,
            "Total Weight (g)":   round(total_weight, 2),
            "Conclusion":         conclusion,
            "Suggested Use":      suggested_use,
        }

    return None