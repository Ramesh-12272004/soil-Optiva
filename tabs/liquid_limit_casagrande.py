import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches

def run():
    st.subheader("Liquid Limit – Casagrande Method (IS 2720 Part 5)")

    # =========================================================
    # PROCEDURE
    # =========================================================
    procedure_text = """
🎯 Objective:
To determine the Liquid Limit (LL) of fine-grained soil using the Casagrande percussion cup apparatus as per IS 2720 (Part 5) – 1985.

🧪 Apparatus Required:
- Casagrande liquid limit device with standard grooving tool
- Porcelain evaporating dish
- Spatula / palette knife
- Weighing balance (accuracy 0.01 g)
- Moisture content cans (at least 4)
- Oven (105°C – 110°C)
- IS sieve 425 μm (No. 40)
- Wash bottle with distilled water
- Desiccator

📝 Theory:
The Liquid Limit is defined as the water content (expressed as a percentage of the dry weight of soil) at which a soil changes from the liquid state to the plastic state. At the LL, a standard groove cut in soil paste in the Casagrande cup just closes over a length of 12.5 mm at exactly 25 blows.

📋 Step-by-Step Procedure:
1. Take about 120 g of air-dried soil passing through the 425 μm IS sieve.
2. Place the soil in a porcelain dish. Add distilled water gradually and mix thoroughly with a spatula to form a uniform, homogeneous paste.
3. The initial water content should be set so the groove closes at approximately 30–35 blows (start from the higher moisture end).
4. Adjust and fix the height of fall of the Casagrande cup to exactly 10 mm using the gauge on the device.
5. Place a portion of the prepared soil paste into the cup, spreading it to a maximum depth of 10 mm at the deepest point.
6. Using the standard ASTM/IS grooving tool, cut a clean groove centrally through the soil, dividing it into two halves. The groove should have dimensions: 2 mm wide at bottom, 11 mm at top, and 8 mm deep.
7. Rotate the crank at 2 revolutions per second. Count and record the number of blows required to close the bottom of the groove over a length of 12.5 mm.
8. Remove a representative soil sample from the closed part of the groove and immediately determine its moisture content.
9. Remove the remaining soil from the cup, mix it with the soil in the dish, and adjust the moisture content (either add water or allow to dry slightly).
10. Repeat steps 5–9 for at least 3 more trials, covering the range of 10 to 40 blows.
11. Record blow counts and moisture contents for a minimum of 4 trials.

⚠️ Important Precautions:
- The groove must be cut in a single stroke to avoid tearing.
- Ensure the cup drops freely from exactly 10 mm.
- Readings close to 25 blows are the most accurate for LL determination.
- Soil must be uniform and free of air pockets before each trial.
- Use distilled water only.
- Do not interpolate beyond the range of recorded data.
"""

    # =========================================================
    # FORMULAS
    # =========================================================
    formulas_text = """
📐 Key Formulas:

1. Moisture Content (w%) for each trial:
   w (%) = [(W₂ − W₃) / (W₃ − W₁)] × 100

   Where:
   • W₁ = Mass of empty moisture can (g)
   • W₂ = Mass of can + wet soil (g)
   • W₃ = Mass of can + dry soil (g)

2. Flow Curve (Semi-log graph):
   Plot moisture content (%) on Y-axis vs.
   Number of blows (N) on X-axis (log scale).
   A best-fit straight line is drawn through the data points.

3. Liquid Limit (LL):
   LL = moisture content (%) read from the flow curve at N = 25 blows.
   Using regression:
   LL = a × log₁₀(25) + b
   Where a, b are regression coefficients from linear fit on log(N) vs. w(%).

4. Flow Index (Iᶠ):
   Iᶠ = (w₁ − w₂) / [log₁₀(N₂/N₁)]
   Represents the slope of the flow curve. Higher Iᶠ → more sensitive soil.

5. Toughness Index (Iᵤ):
   Iᵤ = PI / Iᶠ
   Where PI = Plasticity Index = LL − PL
   (PL = Plastic Limit from separate test)
"""

    st.markdown("## 📘 Test Procedure")
    with st.expander("📖 Click to read full procedure", expanded=False):
        st.markdown(procedure_text)

    st.markdown("## 📐 Formulas Used")
    with st.expander("🔢 Click to view formulas", expanded=False):
        st.markdown(formulas_text)

    # =========================================================
    # SAFE MOISTURE CONTENT CALCULATION
    # =========================================================
    def calculate_moisture_content(w1, w2, w3):
        if not (w2 >= w3 >= w1):
            return np.nan
        denominator = w3 - w1
        if denominator == 0:
            return np.nan
        return ((w2 - w3) / denominator) * 100

    # =========================================================
    # INPUT SECTION
    # =========================================================
    st.markdown("## 📥 Enter Trial Data")
    num_samples = st.number_input("Number of Trials", min_value=2, max_value=10, value=4, step=1)

    if "ll_casagrande_trials" not in st.session_state:
        st.session_state.ll_casagrande_trials = {}

    # Initialize missing trial keys
    for i in range(int(num_samples)):
        key = f"trial_{i+1}"
        if key not in st.session_state.ll_casagrande_trials:
            st.session_state.ll_casagrande_trials[key] = {
                "Number of Blows": 0.0,
                "W1 – Empty Can (g)": 0.0,
                "W2 – Wet Soil+Can (g)": 0.0,
                "W3 – Dry Soil+Can (g)": 0.0,
                "Moisture Content (%)": np.nan,
            }

    trial_data = st.session_state.ll_casagrande_trials

    for i in range(int(num_samples)):
        key = f"trial_{i+1}"
        st.markdown(f"### 🔬 Trial {i+1}")
        c1, c2, c3, c4 = st.columns(4)

        trial_data[key]["Number of Blows"] = c1.number_input(
            "Number of Blows (N)",
            min_value=0.0, step=1.0, value=float(trial_data[key]["Number of Blows"]),
            key=f"ll_blows_{i}"
        )
        trial_data[key]["W1 – Empty Can (g)"] = c2.number_input(
            "W1 – Empty Can (g)",
            min_value=0.0, step=0.01, value=float(trial_data[key]["W1 – Empty Can (g)"]),
            key=f"ll_w1_{i}"
        )
        trial_data[key]["W2 – Wet Soil+Can (g)"] = c3.number_input(
            "W2 – Wet Soil+Can (g)",
            min_value=0.0, step=0.01, value=float(trial_data[key]["W2 – Wet Soil+Can (g)"]),
            key=f"ll_w2_{i}"
        )
        trial_data[key]["W3 – Dry Soil+Can (g)"] = c4.number_input(
            "W3 – Dry Soil+Can (g)",
            min_value=0.0, step=0.01, value=float(trial_data[key]["W3 – Dry Soil+Can (g)"]),
            key=f"ll_w3_{i}"
        )

        mc = calculate_moisture_content(
            trial_data[key]["W1 – Empty Can (g)"],
            trial_data[key]["W2 – Wet Soil+Can (g)"],
            trial_data[key]["W3 – Dry Soil+Can (g)"],
        )
        trial_data[key]["Moisture Content (%)"] = mc

        if np.isnan(mc):
            st.warning(f"Trial {i+1}: Invalid weights — check that W2 ≥ W3 ≥ W1 and W3 ≠ W1.")
        else:
            st.success(f"✅ Trial {i+1} Moisture Content = **{mc:.2f}%**")

    st.session_state.ll_casagrande_trials = trial_data

    # =========================================================
    # CALCULATE
    # =========================================================
    if st.button("🔍 Calculate Liquid Limit"):

        df = pd.DataFrame.from_dict(trial_data, orient="index").reset_index(drop=True)
        df.rename(columns={
            "Number of Blows": "Number of Blows (N)",
            "W1 – Empty Can (g)": "W1 – Empty Can (g)",
            "W2 – Wet Soil+Can (g)": "W2 – Wet Soil+Can (g)",
            "W3 – Dry Soil+Can (g)": "W3 – Dry Soil+Can (g)",
            "Moisture Content (%)": "Moisture Content (%)",
        }, inplace=True)

        df_valid = df[
            (df["Number of Blows (N)"] > 0) &
            (~df["Moisture Content (%)"].isna())
        ].copy()

        if len(df_valid) < 2:
            st.error("❌ At least 2 valid trials required. Check inputs.")
            return None

        # Round for display
        df_display = df_valid.copy()
        df_display["Moisture Content (%)"] = df_display["Moisture Content (%)"].round(2)

        st.markdown("### 📊 Trial Results Table")
        st.dataframe(df_display, use_container_width=True)

        # ---- Regression ----
        log_n = np.log10(df_valid["Number of Blows (N)"].values.astype(float))
        mc    = df_valid["Moisture Content (%)"].values.astype(float)

        try:
            coeffs = np.polyfit(log_n, mc, 1)
            a, b   = coeffs
        except np.linalg.LinAlgError:
            st.error("Curve fitting failed — check input data.")
            return None

        LL = a * np.log10(25) + b

        # Flow Index
        n_min, n_max = df_valid["Number of Blows (N)"].min(), df_valid["Number of Blows (N)"].max()
        w_at_nmin    = a * np.log10(n_min) + b
        w_at_nmax    = a * np.log10(n_max) + b
        flow_index   = (w_at_nmin - w_at_nmax) / np.log10(n_max / n_min) if n_max != n_min else None

        # ---- FLOW CURVE PLOT ----
        st.markdown("### 📈 Flow Curve (Semi-Log)")
        fig, ax = plt.subplots(figsize=(8, 5))

        n_range   = np.linspace(max(1, n_min * 0.85), n_max * 1.15, 200)
        mc_fitted = a * np.log10(n_range) + b

        ax.semilogx(df_valid["Number of Blows (N)"], mc, "o",
                    color="#0a68cc", markersize=9, label="Trial Data", zorder=5)
        ax.semilogx(n_range, mc_fitted, "-",
                    color="#e05c00", linewidth=2, label="Best-Fit Line")
        ax.axvline(x=25, color="gray", linestyle="--", linewidth=1.2, alpha=0.7, label="N = 25")
        ax.axhline(y=LL, color="green", linestyle="--", linewidth=1.2, alpha=0.7,
                   label=f"LL = {LL:.1f}%")
        ax.plot(25, LL, "r*", markersize=14, label=f"LL Point ({LL:.2f}%)", zorder=6)

        ax.set_xlabel("Number of Blows (N) – Log Scale", fontsize=11)
        ax.set_ylabel("Moisture Content (%)", fontsize=11)
        ax.set_title("Flow Curve – Casagrande Liquid Limit Test", fontsize=13, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(True, which="both", linestyle="--", alpha=0.5)
        ax.set_xlim([max(1, n_min * 0.75), n_max * 1.4])

        st.pyplot(fig)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        # ---- RESULTS ----
        st.markdown("### ✅ Computed Results")
        rc1, rc2, rc3 = st.columns(3)
        rc1.metric("Liquid Limit (LL)", f"{LL:.2f}%")
        rc2.metric("Flow Index (Iᶠ)", f"{flow_index:.3f}" if flow_index else "N/A")
        rc3.metric("Valid Trials Used", str(len(df_valid)))

        # Classification
        if LL < 35:
            soil_class = "Low plasticity (L) — Sandy / Silty soil"
        elif LL < 50:
            soil_class = "Medium plasticity (I) — Silty clay"
        else:
            soil_class = "High plasticity (H) — Expansive / Fat clay"

        st.info(f"🏷️ **Soil Classification (based on LL):** {soil_class}")

        conclusion = (
            f"The Liquid Limit of the given soil sample is **{LL:.2f}%** "
            f"(determined at 25 blows from the flow curve). "
            f"The soil exhibits {soil_class.split('—')[0].strip().lower()} characteristics. "
            f"{'This indicates the soil has high compressibility and poor shear strength; suitable stabilization should be considered before use in construction.' if LL >= 50 else 'The soil is moderately suitable for foundation use subject to further Atterberg limit analysis.'}"
        )
        st.markdown("### 📝 Conclusion")
        st.write(conclusion)

        # =========================================================
        # WORD REPORT
        # =========================================================
        doc = Document()
        doc.add_heading("ANITS – Liquid Limit Test Report (Casagrande Method)", 0)
        doc.add_heading("Reference Standard: IS 2720 (Part 5) – 1985", level=2)

        doc.add_heading("Test Procedure", level=1)
        for ln in procedure_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())

        doc.add_heading("Formulas Used", level=1)
        for ln in formulas_text.strip().split("\n"):
            doc.add_paragraph(ln.strip())

        doc.add_page_break()

        doc.add_heading("Observation Table", level=1)
        tbl = doc.add_table(rows=1, cols=len(df_display.columns))
        tbl.style = "Table Grid"
        for idx, col in enumerate(df_display.columns):
            tbl.rows[0].cells[idx].text = str(col)
        for _, row in df_display.iterrows():
            cells = tbl.add_row().cells
            for idx, val in enumerate(row):
                cells[idx].text = str(round(val, 3) if isinstance(val, float) else val)

        doc.add_heading("Flow Curve", level=1)
        img_buf.seek(0)
        doc.add_picture(img_buf, width=Inches(5.5))

        doc.add_heading("Results", level=1)
        doc.add_paragraph(f"Liquid Limit (LL) = {LL:.2f}%")
        if flow_index:
            doc.add_paragraph(f"Flow Index (Iᶠ) = {flow_index:.3f}")
        doc.add_paragraph(f"Soil Classification (LL-based): {soil_class}")

        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(conclusion)

        word_buf = BytesIO()
        doc.save(word_buf)
        word_buf.seek(0)

        st.download_button(
            label="⬇️ Download Word Report",
            data=word_buf,
            file_name="Liquid_Limit_Casagrande_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # =========================================================
        # RETURN RESULT DICT
        # =========================================================
        return {
            "procedure":             procedure_text,
            "formulas":              formulas_text,
            "data":                  df_display,
            "graph":                 img_buf,
            "Liquid Limit LL (%)":   round(LL, 2),
            "Flow Index (If)":       round(flow_index, 3) if flow_index else None,
            "Valid Trials":          len(df_valid),
            "Soil Classification":   soil_class,
            "Conclusion":            conclusion,
        }

    return None