import streamlit as st
import math
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the Optimum Moisture Content (OMC) and Maximum Dry Density
  (MDD) of soil using the Light Compaction (Standard Proctor) test as per
  IS 2720 (Part 7) – 1980.

Apparatus Required:
  - IS Light compaction mould (internal diameter 100 mm, effective height 127.3 mm, volume ≈ 1000 cm³)
  - Standard rammer (2.5 kg, free-fall height 300 mm)
  - Steel straight edge
  - Weighing balance (accuracy 1 g)
  - Moisture content cans
  - Oven (105 °C – 110 °C)
  - IS sieve 20 mm (to remove oversized particles)

Theory:
  Compaction is the process of densifying soil by expelling air from
  the voids under applied mechanical energy. At low moisture content the
  soil is stiff and resists compaction; at high moisture content the water
  occupies space needed by solids. At the OMC, water provides lubrication
  without filling voids excessively, yielding the MDD.

  The Zero-Air-Voids (ZAV) line represents theoretical 100 % saturation
  and always lies above the compaction curve.

Step-by-Step Procedure:
  1. Pass soil through a 20 mm IS sieve; record mass of oversize.
  2. Take ≈ 3 kg of air-dried soil; mix with calculated water to reach
     the first target moisture content (usually 2–4% below expected OMC).
  3. Compact in 3 equal layers, 25 blows per layer with the standard rammer.
  4. Trim flush with the top of the mould; weigh mould + soil. Record W5.
  5. Extract a sample from the top and bottom thirds for moisture content.
     Record W1 (can), W2 (can + wet soil), W3 (can + dry soil).
  6. Extrude, break up, and add more water; repeat for 5–6 points covering
     ±3–4% either side of estimated OMC.
  7. Plot dry density vs water content; the peak is the MDD at the OMC.

Precautions:
  - Ensure uniform mixing; no lumps before compaction.
  - Always compact in exactly 3 layers with exactly 25 blows each.
  - Take moisture samples immediately after compaction.
  - Do not reuse previously compacted soil.
"""

FORMULAS = """
Volume of Mould:
  V = (π / 4) × D² × H    [cm³]

Water Content (w%):
  w (%) = [(W2 – W3) / (W3 – W1)] × 100

Wet (Bulk) Density:
  ρ_wet = (W5 – W4) / V    [g/cm³]

Dry Density:
  ρ_dry = ρ_wet / (1 + w / 100)    [g/cm³]

Where:
  D  = Internal diameter of mould (cm)
  H  = Internal height of mould (cm)
  W1 = Mass of empty moisture can (g)
  W2 = Mass of can + wet soil (g)
  W3 = Mass of can + dry soil (g)
  W4 = Mass of empty mould + base plate (g)
  W5 = Mass of mould + compacted soil + base plate (g)

Zero Air Voids (ZAV) line:
  ρ_zav = G_s × ρ_w / (1 + G_s × w / 100)
  (G_s ≈ 2.68 for most soils; ρ_w = 1 g/cm³)
"""


def _generate_report(df, mdd, omc, volume, d, h, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("Light Compaction Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 7) – 1980")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the Optimum Moisture Content (OMC) and Maximum Dry Density (MDD) "
        "of the soil using the Light (Standard Proctor) compaction test."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph(f"Mould Diameter D = {d:.2f} cm")
    doc.add_paragraph(f"Mould Height  H = {h:.2f} cm")
    doc.add_paragraph(f"Mould Volume  V = {volume:.2f} cm³")

    doc.add_heading("Compaction Data Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.3f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Maximum Dry Density (MDD) = {mdd:.3f} g/cm³")
    doc.add_paragraph(f"Optimum Moisture Content (OMC) = {omc:.2f} %")

    doc.add_heading("Compaction Curve", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The MDD is {mdd:.3f} g/cm³ at an OMC of {omc:.2f}%. "
        f"{'Good compaction characteristics.' if mdd >= 1.7 else 'Moderate compaction characteristics; stabilisation may be considered.'}"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("🔨 Light Compaction Test – Standard Proctor (IS 2720 Part 7 : 1980)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "lc_res" not in st.session_state:
        st.session_state.lc_res = None

    # ── Mould dimensions ──
    st.markdown("### 📏 Mould Dimensions")
    c1, c2 = st.columns(2)
    d = c1.number_input("Diameter D (cm)", value=10.0, min_value=1.0, format="%.2f", key="lc_d")
    h = c2.number_input("Height H (cm)",   value=12.7, min_value=1.0, format="%.2f", key="lc_h")
    w4 = st.number_input("W4 – Mass of Empty Mould + Base (g)", value=0.0, min_value=0.0, format="%.2f", key="lc_w4")

    volume = math.pi / 4 * d ** 2 * h
    st.info(f"📐 Volume V = **{volume:.2f} cm³**")

    # ── Number of trials ──
    num_trials = st.number_input("Number of Compaction Points", min_value=3, max_value=10, value=5, step=1, key="lc_nt")

    if "lc_inputs" not in st.session_state or len(st.session_state.lc_inputs) != num_trials:
        st.session_state.lc_inputs = [
            {"w1": 0.0, "w2": 0.0, "w3": 0.0, "w5": 0.0}
            for _ in range(num_trials)
        ]

    # ── Trial data ──
    st.markdown("### 📋 Enter Compaction Trial Data")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3, c4 = st.columns(4)
        inp = st.session_state.lc_inputs[i]
        inp["w1"] = c1.number_input(f"W1 – Empty Can (g)",       value=inp["w1"], min_value=0.0, format="%.3f", key=f"lc_w1_{i}")
        inp["w2"] = c2.number_input(f"W2 – Can + Wet Soil (g)",  value=inp["w2"], min_value=0.0, format="%.3f", key=f"lc_w2_{i}")
        inp["w3"] = c3.number_input(f"W3 – Can + Dry Soil (g)",  value=inp["w3"], min_value=0.0, format="%.3f", key=f"lc_w3_{i}")
        inp["w5"] = c4.number_input(f"W5 – Mould + Comp. Soil + Base (g)", value=inp["w5"], min_value=0.0, format="%.3f", key=f"lc_w5_{i}")

    if st.button("🔄 Reset"):
        st.session_state.lc_inputs = [{"w1": 0.0, "w2": 0.0, "w3": 0.0, "w5": 0.0} for _ in range(num_trials)]
        st.session_state.lc_res = None
        st.rerun()

    if st.button("📊 Calculate"):
        rows = []
        for i, inp in enumerate(st.session_state.lc_inputs[:num_trials]):
            w1, w2, w3, w5 = inp["w1"], inp["w2"], inp["w3"], inp["w5"]
            if w2 > w3 > w1 > 0 and w5 > w4 > 0:
                Wd = w3 - w1
                wc = (w2 - w3) / Wd * 100
                rho_wet = (w5 - w4) / volume
                rho_dry = rho_wet / (1 + wc / 100)
                rows.append({
                    "Trial":              i + 1,
                    "Water Content (%)":  round(wc, 2),
                    "Wet Density (g/cm³)": round(rho_wet, 3),
                    "Dry Density (g/cm³)": round(rho_dry, 3),
                })
            else:
                rows.append({
                    "Trial": i + 1,
                    "Water Content (%)": None,
                    "Wet Density (g/cm³)": None,
                    "Dry Density (g/cm³)": None,
                })

        df = pd.DataFrame(rows).dropna()
        if df.empty or len(df) < 2:
            st.error("At least 2 valid trials required. Check that W2 > W3 > W1 and W5 > W4.")
            return None

        df = df.sort_values("Water Content (%)")
        mdd = float(df["Dry Density (g/cm³)"].max())
        omc = float(df.loc[df["Dry Density (g/cm³)"].idxmax(), "Water Content (%)"])

        # ZAV line
        Gs = 2.68
        wc_range = range(int(df["Water Content (%)"].min()), int(df["Water Content (%)"].max()) + 5)
        zav_wc  = [w for w in wc_range]
        zav_dd  = [Gs / (1 + Gs * w / 100) for w in wc_range]

        # Plot
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(df["Water Content (%)"], df["Dry Density (g/cm³)"],
                marker="o", color="#0a68cc", linewidth=2, label="Compaction Curve")
        ax.plot(zav_wc, zav_dd, "--", color="#888", linewidth=1.2, label="ZAV Line (Gs=2.68)")
        ax.plot(omc, mdd, "r*", markersize=14, label=f"MDD={mdd:.3f} g/cm³ @ OMC={omc:.1f}%")
        ax.axvline(omc, linestyle=":", color="red", alpha=0.5)
        ax.axhline(mdd, linestyle=":", color="red", alpha=0.5)
        ax.set_xlabel("Water Content (%)", fontsize=11)
        ax.set_ylabel("Dry Density (g/cm³)", fontsize=11)
        ax.set_title("Compaction Curve (Light Compaction)", fontsize=13, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(True, linestyle="--", alpha=0.4)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.session_state.lc_res = {
            "df": df, "mdd": mdd, "omc": omc, "volume": volume,
            "d": d, "h": h, "img_buf": img_buf
        }

    if st.session_state.lc_res is not None:
        res = st.session_state.lc_res
        st.markdown("### 📋 Compaction Data")
        st.dataframe(res["df"], use_container_width=True)
        c1, c2 = st.columns(2)
        c1.metric("MDD (g/cm³)", f"{res['mdd']:.3f}")
        c2.metric("OMC (%)",     f"{res['omc']:.2f}")
        st.image(res["img_buf"])

        report_buf = _generate_report(
            res["df"], res["mdd"], res["omc"],
            res["volume"], res["d"], res["h"],
            res["img_buf"], PROCEDURE, FORMULAS
        )
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Light_Compaction_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":                   PROCEDURE,
            "formulas":                    FORMULAS,
            "data":                        res["df"],
            "graph":                       res["img_buf"],
            "Maximum Dry Density MDD (g/cm3)": round(res["mdd"], 3),
            "Optimum Moisture Content OMC (%)": round(res["omc"], 2),
        }

    return None