import streamlit as st
import math
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the Unconfined Compressive Strength (UCS) and undrained
  shear strength of cohesive soils as per IS 2720 (Part 10) – 1991.

Apparatus Required:
  - UCS testing machine (strain-controlled)
  - Proving ring with dial gauge
  - Vertical deformation (strain) dial gauge
  - Steel scale and vernier calliper
  - Stopwatch
  - Sample extruder
  - Soil trimming equipment (wire saw, palette knife)

Theory:
  In the UCS test the soil specimen (cylindrical, H/D ≈ 2) is sheared under
  zero confining pressure. Since it is an undrained test on saturated clay,
  the total stress friction angle ϕ = 0 and the undrained shear strength:
  
    S_u = UCS / 2

  UCS is sensitive to sample disturbance. Typical ranges:
    UCS < 25 kPa   → Very soft clay
    25 – 50 kPa    → Soft clay
    50 – 100 kPa   → Medium clay
    100 – 200 kPa  → Stiff clay
    > 200 kPa      → Very stiff / hard clay

Step-by-Step Procedure:
  1. Measure the initial diameter (D, mm) and length (L, mm) of each specimen.
  2. Place the specimen centrally in the loading frame; zero all gauges.
  3. Apply load at a constant strain rate of 0.5–2% per minute.
  4. Record the proving ring reading at regular strain intervals (0.5% or 1%).
  5. Stop when the load drops sharply (failure) or at 20% axial strain.
  6. Sketch or photograph the failure plane.
  7. Repeat for additional specimens (min 3 for a reliable average).

Precautions:
  - Do not use specimens with obvious cracks or significant disturbance.
  - Keep end surfaces parallel and flat.
  - Apply load centrally to avoid eccentric loading.
  - Correct the cross-sectional area for axial strain:
    A_corrected = A₀ / (1 – ε)  where ε = axial strain (decimal).
"""

FORMULAS = """
Initial Cross-Sectional Area:
  A₀ = π × (D / 2)²    [mm²]

Failure Load:
  P = Proving Ring Constant (N/div) × Ring Reading (div)    [N]

Axial Strain at Failure:
  ε = ΔL / L₀

Corrected Area (area correction):
  A_c = A₀ / (1 – ε)    [mm²]

Unconfined Compressive Strength:
  UCS = P / A_c × 1000    [kPa]
  (factor 1000 converts N/mm² → kPa; note 1 N/mm² = 1000 kPa)

Undrained Shear Strength:
  S_u = UCS / 2    [kPa]

Where:
  D  = Diameter of specimen (mm)
  L₀ = Initial length of specimen (mm)
  ΔL = Change in length at failure (mm)
"""


def _generate_report(df, mean_ucs, mean_su, procedure, formulas):
    doc = Document()
    doc.add_heading("Unconfined Compressive Strength (UCS) Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 10) – 1991")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the UCS and undrained shear strength Su of cohesive soil."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Results Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.2f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Average UCS = {mean_ucs:.2f} kPa")
    doc.add_paragraph(f"Average Undrained Shear Strength Su = {mean_su:.2f} kPa")

    # Chart
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(df["Trial"], df["UCS (kPa)"], color="#4a90d9", edgecolor="white", linewidth=1.2)
    ax.axhline(mean_ucs, color="#e05c00", linestyle="--", linewidth=1.5,
               label=f"Avg UCS = {mean_ucs:.2f} kPa")
    ax.set_xlabel("Trial")
    ax.set_ylabel("UCS (kPa)")
    ax.set_title("UCS per Specimen")
    ax.legend()
    ax.grid(True, axis="y", linestyle="--", alpha=0.5)

    img_buf = BytesIO()
    fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
    img_buf.seek(0)
    plt.close(fig)
    doc.add_heading("UCS Chart", 1)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Conclusion", 1)
    if mean_ucs < 25:
        consistency = "Very Soft Clay"
    elif mean_ucs < 50:
        consistency = "Soft Clay"
    elif mean_ucs < 100:
        consistency = "Medium Clay"
    elif mean_ucs < 200:
        consistency = "Stiff Clay"
    else:
        consistency = "Very Stiff / Hard Clay"
    doc.add_paragraph(
        f"Average UCS = {mean_ucs:.2f} kPa and Su = {mean_su:.2f} kPa. "
        f"Soil consistency: {consistency}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, img_buf


def run():
    st.subheader("🔬 Unconfined Compressive Strength (UCS) Test (IS 2720 Part 10 : 1991)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "ucs_res" not in st.session_state:
        st.session_state.ucs_res = None

    num_trials = st.number_input("Number of Specimens", min_value=1, max_value=10, value=3, step=1, key="ucs_nt")

    if "ucs_inputs" not in st.session_state or len(st.session_state.ucs_inputs) != num_trials:
        st.session_state.ucs_inputs = [
            {"D": 0.0, "L": 0.0, "dL": 0.0, "k": 0.0, "r": 0.0}
            for _ in range(num_trials)
        ]

    st.markdown("### 📋 Enter Specimen Data")
    for i in range(num_trials):
        st.markdown(f"#### Specimen {i + 1}")
        c1, c2, c3, c4, c5 = st.columns(5)
        inp = st.session_state.ucs_inputs[i]
        inp["D"]  = c1.number_input(f"D (mm)",  value=inp["D"],  min_value=0.0, format="%.2f", key=f"ucs_D_{i}")
        inp["L"]  = c2.number_input(f"L₀ (mm)", value=inp["L"],  min_value=0.0, format="%.2f", key=f"ucs_L_{i}")
        inp["dL"] = c3.number_input(f"ΔL at failure (mm)", value=inp["dL"], min_value=0.0, format="%.3f", key=f"ucs_dL_{i}")
        inp["k"]  = c4.number_input(f"PR Constant (N/div)", value=inp["k"], min_value=0.0, format="%.4f", key=f"ucs_k_{i}")
        inp["r"]  = c5.number_input(f"PR Reading (div)", value=inp["r"], min_value=0.0, format="%.2f", key=f"ucs_r_{i}")

    if st.button("🔄 Reset"):
        st.session_state.ucs_inputs = [{"D": 0.0, "L": 0.0, "dL": 0.0, "k": 0.0, "r": 0.0} for _ in range(num_trials)]
        st.session_state.ucs_res = None
        st.rerun()

    if st.button("📊 Calculate"):
        rows = []
        for i, inp in enumerate(st.session_state.ucs_inputs[:num_trials]):
            D, L, dL, k, r = inp["D"], inp["L"], inp["dL"], inp["k"], inp["r"]
            if D > 0 and L > 0 and k > 0 and r > 0:
                A0  = math.pi * (D / 2) ** 2
                eps = dL / L if L > 0 else 0.0
                Ac  = A0 / (1 - eps) if eps < 1 else A0
                P   = k * r
                ucs = (P / Ac) * 1000   # N/mm² → kPa (since 1 N/mm² = 1 MPa ≠ kPa)
                # Correct: P[N] / Ac[mm²] = N/mm² = MPa; ×1000 = kPa
                # Actually 1 N/mm² = 1 MPa = 1000 kPa — let's keep as N/mm² × 1000 = kPa
                # Simpler: P(N) / Ac(mm²) = MPa; UCS(kPa) = MPa × 1000
                su  = ucs / 2
                rows.append({
                    "Trial":                i + 1,
                    "D (mm)":               D,
                    "L₀ (mm)":              L,
                    "ΔL (mm)":              dL,
                    "A₀ (mm²)":             round(A0, 2),
                    "ε":                    round(eps, 4),
                    "A_c (mm²)":            round(Ac, 2),
                    "P (N)":                round(P, 2),
                    "UCS (kPa)":            round(ucs, 2),
                    "Su (kPa)":             round(su, 2),
                })

        if not rows:
            st.error("No valid specimens. Check that D, L, k, and r are all > 0.")
            return None

        df = pd.DataFrame(rows)
        mean_ucs = float(df["UCS (kPa)"].mean())
        mean_su  = float(df["Su (kPa)"].mean())

        st.markdown("### 📋 Results Table")
        st.dataframe(df, use_container_width=True)

        c1, c2 = st.columns(2)
        c1.metric("Average UCS (kPa)", f"{mean_ucs:.2f}")
        c2.metric("Average Su (kPa)",  f"{mean_su:.2f}")

        if mean_ucs < 25:
            consistency = "Very Soft Clay"
        elif mean_ucs < 50:
            consistency = "Soft Clay"
        elif mean_ucs < 100:
            consistency = "Medium Clay"
        elif mean_ucs < 200:
            consistency = "Stiff Clay"
        else:
            consistency = "Very Stiff / Hard Clay"
        st.info(f"🏷️ Soil Consistency: **{consistency}**")

        report_buf, img_buf = _generate_report(df, mean_ucs, mean_su, PROCEDURE, FORMULAS)
        st.image(img_buf)

        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="UCS_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.session_state.ucs_res = {
            "df": df, "mean_ucs": mean_ucs, "mean_su": mean_su,
            "consistency": consistency, "img_buf": img_buf,
        }

        return {
            "procedure":            PROCEDURE,
            "formulas":             FORMULAS,
            "data":                 df,
            "graph":                img_buf,
            "Average UCS (kPa)":    round(mean_ucs, 2),
            "Average Su (kPa)":     round(mean_su, 2),
            "Consistency":          consistency,
        }

    return None