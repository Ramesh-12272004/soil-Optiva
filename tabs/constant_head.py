import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


PROCEDURE = """
Objective:
  To determine the coefficient of permeability (k) of coarse-grained
  soils using the constant head apparatus as per IS 2720 (Part 36) – 1987.

Apparatus Required:
  - Constant head permeameter with overflow outlet
  - Graduated cylinder (measuring jar)
  - Stopwatch
  - Steel scale / vernier calliper
  - Weighing balance
  - Filter paper, porous stones

Theory:
  In a constant head test the hydraulic head (h) driving flow through
  the specimen is kept constant. By measuring the volume of water Q
  that seeps through the specimen of length L and cross-section A in
  time t, Darcy's law gives the coefficient of permeability k.

  Typical k ranges (IS classification):
    k > 10⁻¹ cm/s   → Gravel
    10⁻² – 10⁻¹     → Coarse Sand
    10⁻³ – 10⁻²     → Medium Sand
    10⁻⁴ – 10⁻³     → Fine Sand
    10⁻⁶ – 10⁻⁴     → Silt
    < 10⁻⁶ cm/s     → Clay

Step-by-Step Procedure:
  1. Measure and record the internal dimensions of the permeameter mould.
  2. Compact or place the soil specimen in the mould; fix porous stones.
  3. Connect the inlet to the constant-head reservoir; allow saturation
     (upward flow for at least 30 min).
  4. Adjust the overflow to set a constant head difference h (cm).
  5. Allow steady flow; note the time t (s) for collecting volume Q (cm³).
  6. Repeat 3 times at the same head to verify consistency.
  7. If desired, vary h and repeat for a range of heads.

Precautions:
  - Ensure full saturation before starting measurements.
  - Measure h between the inlet and outlet water surfaces, not piezometers.
  - Record Q accurately using a measuring cylinder.
  - Avoid entrapped air by de-airing water or using upward flow.
"""

FORMULAS = """
Coefficient of Permeability (k):
  k = (Q × L) / (A × h × t)    [cm/s]

Where:
  Q = Volume of water collected (cm³)
  L = Length / height of soil specimen (cm)
  A = Cross-sectional area of specimen (cm²)
  h = Constant head difference (cm)
  t = Time of collection (s)

Conversion:
  k (m/s) = k (cm/s) × 0.01
"""


def _classify(k_cms):
    if k_cms > 1e-1:
        return "Gravel"
    elif k_cms > 1e-2:
        return "Coarse Sand"
    elif k_cms > 1e-3:
        return "Medium Sand"
    elif k_cms > 1e-4:
        return "Fine Sand"
    elif k_cms > 1e-6:
        return "Silt"
    else:
        return "Clay"


def _generate_report(df, avg_k, soil_type, procedure, formulas):
    doc = Document()
    doc.add_heading("Constant Head Permeability Test Report", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 36) – 1987")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the coefficient of permeability k of coarse-grained soil "
        "using the constant head method."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Results Table", 1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = f"{val:.5f}" if isinstance(val, float) else str(val)

    doc.add_heading("Summary", 1)
    doc.add_paragraph(f"Average k = {avg_k:.5f} cm/s  ({avg_k * 0.01:.5e} m/s)")
    doc.add_paragraph(f"Soil Classification: {soil_type}")

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The coefficient of permeability is {avg_k:.5f} cm/s, "
        f"indicating the soil is {soil_type}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run():
    st.subheader("💧 Constant Head Permeability Test (IS 2720 Part 36 : 1987)")

    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    if "ch_res" not in st.session_state:
        st.session_state.ch_res = None

    num_trials = st.number_input("Number of Trials", min_value=1, max_value=10, value=3, step=1, key="ch_nt")

    if "ch_inputs" not in st.session_state or len(st.session_state.ch_inputs) != num_trials:
        st.session_state.ch_inputs = [
            {"L": 0.0, "A": 0.0, "h": 0.0, "Q": 0.0, "t": 0.0}
            for _ in range(num_trials)
        ]

    st.markdown("### 📋 Enter Trial Data")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3 = st.columns(3)
        inp = st.session_state.ch_inputs[i]
        inp["L"] = c1.number_input(f"L – Specimen Length (cm)", value=inp["L"], min_value=0.0, format="%.3f", key=f"ch_L_{i}")
        inp["A"] = c1.number_input(f"A – Cross-section (cm²)",  value=inp["A"], min_value=0.0, format="%.3f", key=f"ch_A_{i}")
        inp["h"] = c2.number_input(f"h – Head Difference (cm)", value=inp["h"], min_value=0.0, format="%.3f", key=f"ch_h_{i}")
        inp["Q"] = c2.number_input(f"Q – Volume Collected (cm³)", value=inp["Q"], min_value=0.0, format="%.3f", key=f"ch_Q_{i}")
        inp["t"] = c3.number_input(f"t – Time (s)",              value=inp["t"], min_value=0.0, format="%.3f", key=f"ch_t_{i}")

    if st.button("🔄 Reset"):
        st.session_state.ch_inputs = [{"L": 0.0, "A": 0.0, "h": 0.0, "Q": 0.0, "t": 0.0} for _ in range(num_trials)]
        st.session_state.ch_res = None
        st.rerun()

    if st.button("📊 Calculate"):
        rows = []
        for i, inp in enumerate(st.session_state.ch_inputs[:num_trials]):
            L, A, h, Q, t = inp["L"], inp["A"], inp["h"], inp["Q"], inp["t"]
            k = (Q * L) / (A * h * t) if (A > 0 and h > 0 and t > 0 and L > 0 and Q > 0) else 0.0
            rows.append({
                "Trial":           i + 1,
                "L (cm)":          L,
                "A (cm²)":         A,
                "h (cm)":          h,
                "Q (cm³)":         Q,
                "t (s)":           t,
                "k (cm/s)":        round(k, 6),
                "k (m/s)":         round(k * 0.01, 8),
            })

        df = pd.DataFrame(rows)
        valid_k = df[df["k (cm/s)"] > 0]["k (cm/s)"]

        if valid_k.empty:
            st.error("No valid trials. Check that all inputs are positive.")
            return None

        avg_k    = float(valid_k.mean())
        soil_type = _classify(avg_k)

        # Plot
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.plot(df["Trial"], df["k (cm/s)"], marker="o", color="#0a68cc", linewidth=2)
        ax.axhline(avg_k, color="#e05c00", linestyle="--", linewidth=1.5, label=f"Avg k = {avg_k:.5f}")
        ax.set_xlabel("Trial Number")
        ax.set_ylabel("k (cm/s)")
        ax.set_title("Coefficient of Permeability per Trial")
        ax.legend()
        ax.grid(True, linestyle="--", alpha=0.5)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        st.session_state.ch_res = {
            "df": df, "avg_k": avg_k, "soil_type": soil_type, "img_buf": img_buf
        }

    if st.session_state.ch_res is not None:
        res = st.session_state.ch_res
        st.markdown("### 📋 Results Table")
        st.dataframe(res["df"].style.format(precision=6), use_container_width=True)

        c1, c2 = st.columns(2)
        c1.metric("Average k (cm/s)", f"{res['avg_k']:.5f}")
        c2.metric("Soil Classification", res["soil_type"])
        st.image(res["img_buf"])

        report_buf = _generate_report(res["df"], res["avg_k"], res["soil_type"], PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Constant_Head_Permeability_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":              PROCEDURE,
            "formulas":               FORMULAS,
            "data":                   res["df"],
            "graph":                  res["img_buf"],
            "Average k (cm/s)":       round(res["avg_k"], 6),
            "Soil Classification":    res["soil_type"],
        }

    return None