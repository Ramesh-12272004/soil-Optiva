import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime


# ─────────────────────────────────────────────
# HELPER
# ─────────────────────────────────────────────
def _calc_wc(w1, w2, w3):
    """Return water content (%) or NaN if inputs are invalid."""
    if w2 <= w3 or w3 <= w1:
        return np.nan
    return (w2 - w3) / (w3 - w1) * 100


# ─────────────────────────────────────────────
# WORD REPORT
# ─────────────────────────────────────────────
def _generate_report(df_all, liquid_limit, soil_class, img_buf, procedure, formulas):
    doc = Document()
    doc.add_heading("Liquid Limit Test Report – Cone Penetration Method", 0)
    doc.add_paragraph("Reference Standard: IS 2720 (Part 5) – 1985")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_heading("Objective", 1)
    doc.add_paragraph(
        "To determine the Liquid Limit (LL) of fine-grained soil using the "
        "cone penetration apparatus as per IS 2720 (Part 5) – 1985."
    )

    doc.add_heading("Test Procedure", 1)
    for line in procedure.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Formulas Used", 1)
    for line in formulas.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Observation Table", 1)
    table = doc.add_table(rows=1, cols=len(df_all.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df_all.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df_all.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(round(val, 3) if isinstance(val, float) else val)

    doc.add_heading("Flow Curve", 1)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))

    doc.add_heading("Results", 1)
    doc.add_paragraph(f"Liquid Limit (LL) = {liquid_limit:.2f} %")
    doc.add_paragraph(f"Soil Classification: {soil_class}")

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"The Liquid Limit of the given soil sample is {liquid_limit:.2f}% "
        f"(determined at 20 mm penetration from the flow curve). "
        f"The soil is classified as {soil_class}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def run():
    st.subheader("🔵 Liquid Limit – Cone Penetration Method (IS 2720 Part 5 : 1985)")

    PROCEDURE = """
Objective:
  To determine the Liquid Limit (LL) of fine-grained soil using the cone
  penetration apparatus as per IS 2720 (Part 5) – 1985.

Apparatus Required:
  - Cone penetrometer (cone angle 30°, mass 80 g)
  - Porcelain mixing dish
  - Spatula / palette knife
  - Weighing balance (accuracy 0.01 g)
  - Moisture content cans with lids
  - Oven (105°C – 110°C)
  - IS sieve 425 µm
  - Wash bottle with distilled water

Theory:
  In this method the Liquid Limit is the water content at which a standard
  cone (80 g, 30°) penetrates exactly 20 mm into the soil paste in 5 seconds.
  A best-fit straight line (flow curve) is drawn through at least four
  data points (penetration vs water content) and the LL is read at 20 mm.

Step-by-Step Procedure:
  1. Pass air-dried soil through a 425 µm sieve; collect about 200 g.
  2. Place in a porcelain dish; add distilled water to make a stiff paste.
  3. Fill the penetrometer cup with the paste (no air bubbles); level flush.
  4. Lower the cone to just touch the surface; release for exactly 5 seconds.
  5. Record the penetration depth (mm).
  6. Remove a sample from the penetration zone; determine its water content.
  7. Add more water, re-mix thoroughly, and repeat steps 3–6.
  8. Obtain at least 4 trials covering a penetration range of about 15–25 mm.

Precautions:
  - Ensure the cone is clean and undamaged before each trial.
  - Level the paste carefully; avoid air inclusions.
  - Use distilled water only.
  - Release the cone without any jerk.
"""

    FORMULAS = """
Water Content (w%):
  w (%) = [(W2 - W3) / (W3 - W1)] × 100

  Where:
    W1 = Mass of empty moisture-content container (g)
    W2 = Mass of container + wet soil              (g)
    W3 = Mass of container + oven-dried soil       (g)

Flow Curve (linear regression):
  w = a × penetration + b
  (ordinary least-squares fit through the 4+ trial points)

Liquid Limit:
  LL = water content at penetration = 20 mm
  LL = a × 20 + b
"""

    # ── Procedure / Formula expanders ──
    with st.expander("📘 View Detailed Procedure"):
        st.markdown(PROCEDURE)
    with st.expander("📐 View Formulas"):
        st.markdown(FORMULAS)

    # ── Session state ──
    if "cl_num_trials" not in st.session_state:
        st.session_state.cl_num_trials = 4
    if "cl_inputs" not in st.session_state:
        st.session_state.cl_inputs = [
            {"penetration": 0.0, "w1": 0.0, "w2": 0.0, "w3": 0.0}
            for _ in range(st.session_state.cl_num_trials)
        ]

    num_trials = st.number_input("Number of Trials", min_value=3, max_value=10,
                                 value=st.session_state.cl_num_trials, step=1)

    # Resize input list when trial count changes
    while len(st.session_state.cl_inputs) < num_trials:
        st.session_state.cl_inputs.append({"penetration": 0.0, "w1": 0.0, "w2": 0.0, "w3": 0.0})
    st.session_state.cl_inputs = st.session_state.cl_inputs[:num_trials]
    st.session_state.cl_num_trials = num_trials

    # ── Input fields ──
    st.markdown("### 📋 Enter Trial Data")
    for i in range(num_trials):
        st.markdown(f"#### Trial {i + 1}")
        c1, c2, c3, c4 = st.columns(4)
        inp = st.session_state.cl_inputs[i]

        inp["penetration"] = c1.number_input(
            f"Penetration (mm)", value=inp["penetration"],
            min_value=0.0, format="%.2f", key=f"cl_pen_{i}"
        )
        inp["w1"] = c2.number_input(
            f"W1 – Empty Can (g)", value=inp["w1"],
            min_value=0.0, format="%.3f", key=f"cl_w1_{i}"
        )
        inp["w2"] = c3.number_input(
            f"W2 – Wet Soil+Can (g)", value=inp["w2"],
            min_value=0.0, format="%.3f", key=f"cl_w2_{i}"
        )
        inp["w3"] = c4.number_input(
            f"W3 – Dry Soil+Can (g)", value=inp["w3"],
            min_value=0.0, format="%.3f", key=f"cl_w3_{i}"
        )

        wc = _calc_wc(inp["w1"], inp["w2"], inp["w3"])
        if np.isnan(wc):
            if not (inp["w1"] == inp["w2"] == inp["w3"] == 0.0):
                st.warning(f"Trial {i + 1}: Invalid weights (need W2 > W3 > W1).")
        else:
            st.info(f"Trial {i + 1} Water Content = **{wc:.2f}%**")

    # ── Reset ──
    if st.button("🔄 Reset All Inputs"):
        st.session_state.cl_inputs = [
            {"penetration": 0.0, "w1": 0.0, "w2": 0.0, "w3": 0.0}
            for _ in range(num_trials)
        ]
        st.rerun()

    # ── Calculate ──
    if st.button("📊 Calculate Liquid Limit"):
        rows = []
        for i, inp in enumerate(st.session_state.cl_inputs[:num_trials]):
            wc = _calc_wc(inp["w1"], inp["w2"], inp["w3"])
            rows.append({
                "Trial":              i + 1,
                "Penetration (mm)":   inp["penetration"],
                "W1 (g)":             inp["w1"],
                "W2 (g)":             inp["w2"],
                "W3 (g)":             inp["w3"],
                "Water Content (%)":  round(wc, 2) if not np.isnan(wc) else np.nan,
            })

        df_all = pd.DataFrame(rows)
        df_valid = df_all[
            (df_all["Penetration (mm)"] > 0) &
            df_all["Water Content (%)"].notna()
        ].copy()

        if len(df_valid) < 2:
            st.error("Need at least 2 valid trials with non-zero penetration and correct weights.")
            return None

        if df_valid["Penetration (mm)"].nunique() < 2:
            st.error("Penetration values must be distinct across trials.")
            return None

        x = df_valid["Penetration (mm)"].values.astype(float)
        y = df_valid["Water Content (%)"].values.astype(float)
        coeffs = np.polyfit(x, y, 1)
        poly   = np.poly1d(coeffs)
        ll     = float(poly(20))

        # ── Soil classification ──
        if ll < 35:
            soil_class = "Low Plasticity (L) – Sandy / Silty soil"
        elif ll <= 50:
            soil_class = "Intermediate Plasticity (I) – Silty Clay"
        else:
            soil_class = "High Plasticity (H) – Expansive / Fat Clay"

        # ── Plot ──
        fig, ax = plt.subplots(figsize=(8, 5))
        x_fit = np.linspace(max(0, x.min() - 2), x.max() + 2, 200)
        ax.scatter(x, y, color="#0a68cc", s=70, zorder=5, label="Observed Data")
        ax.plot(x_fit, poly(x_fit), color="#e05c00", linewidth=2, label="Best-Fit Line")
        ax.axvline(20, color="gray", linestyle="--", linewidth=1.2, label="20 mm Penetration")
        ax.axhline(ll, color="green", linestyle="--", linewidth=1.2,
                   label=f"LL = {ll:.2f}%")
        ax.plot(20, ll, "r*", markersize=14, label=f"LL Point ({ll:.2f}%)", zorder=6)
        ax.set_xlabel("Penetration (mm)", fontsize=11)
        ax.set_ylabel("Water Content (%)", fontsize=11)
        ax.set_title("Cone Penetration – Flow Curve", fontsize=13, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(True, linestyle="--", alpha=0.5)
        st.pyplot(fig)

        img_buf = BytesIO()
        fig.savefig(img_buf, format="png", dpi=150, bbox_inches="tight")
        img_buf.seek(0)
        plt.close(fig)

        # ── Display results ──
        st.markdown("### 📊 Observation Table")
        st.dataframe(df_all.round(3), use_container_width=True)

        r1, r2 = st.columns(2)
        r1.metric("Liquid Limit (LL)", f"{ll:.2f}%")
        r2.metric("Soil Classification", soil_class.split("–")[0].strip())
        st.info(f"🏷️ **Classification:** {soil_class}")

        # ── Word report ──
        report_buf = _generate_report(df_all, ll, soil_class, img_buf, PROCEDURE, FORMULAS)
        st.download_button(
            "⬇️ Download Word Report",
            data=report_buf,
            file_name="Cone_Penetration_LL_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "procedure":              PROCEDURE,
            "formulas":               FORMULAS,
            "data":                   df_all,
            "graph":                  img_buf,
            "Liquid Limit LL (%)":    round(ll, 2),
            "Soil Classification":    soil_class,
        }

    return None